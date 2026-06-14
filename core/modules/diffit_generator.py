import json
import re
import random
from io import BytesIO
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from core.ai_service import get_ai_response

def extract_random_slice(text, num_words):
    """
    Extracts a random sequential slice of num_words from the text.
    """
    if not text:
        return ""
    words = text.split()
    if len(words) <= num_words:
        return text
    start_idx = random.randint(0, len(words) - num_words)
    return " ".join(words[start_idx:start_idx + num_words])

def summarise_source_text(
    text,
    method,
    algorithm,
    word_limit,
    ai_provider=None,
    ollama_model=None,
    api_key=None,
    ollama_url=None,
    gemini_model=None
):
    """
    Summarizes source text using traditional extractive methods or AI.
    """
    if not text:
        return ""
    if method == "Traditional Extractive":
        from core.modules.summarisation import summarize_text_extractive
        return summarize_text_extractive(text, language="english", word_limit=word_limit, algorithm=algorithm)
    else:
        from core.modules.summarisation import summarize_text_ai
        model = gemini_model if ai_provider == "Gemini" else ollama_model
        return summarize_text_ai(
            text,
            provider=ai_provider,
            model=model,
            api_key=api_key,
            word_limit=word_limit
        )

def _clean_and_parse_json(response, target_key):
    """Safely extracts and parses JSON for a specific key, handling fallback parsing on LLM errors."""
    if not response:
        return None
    try:
        clean = response.strip()
        start = clean.find('{')
        end = clean.rfind('}')
        if start != -1 and end != -1 and end > start:
            clean = clean[start:end+1]
        # Clean trailing commas
        clean = re.sub(r',\s*([\]}])', r'\1', clean)
        parsed = json.loads(clean)
        return parsed.get(target_key)
    except Exception as e:
        print(f"[WARN] JSON loads failed for {target_key}: {e}. Raw response: {response}")
        # Secondary fallback: try to find key using regex or basic formatting if loads crashed
        if target_key == "adapted_passage":
            match = re.search(r'"adapted_passage"\s*:\s*"([^"]+)"', response)
            if match:
                return match.group(1)
        return None

def generate_diffit_resource(
    source_type,
    source_val,
    grade_level,
    language,
    selected_sections,
    ai_provider,
    ollama_model,
    api_key,
    ollama_url,
    gemini_model,
    status_callback=None
):
    """
    Generates a Diffit-like educational resource sequentially (multi-step pipeline)
    to avoid HTTP read timeout issues on local LLMs.
    """
    def log_status(msg):
        if status_callback:
            status_callback(msg)
        print(msg)

    data = {
        "adapted_passage": "",
        "summary": [],
        "vocabulary": [],
        "mcqs": [],
        "short_answers": [],
        "blooms_prompts": []
    }

    # Step 1: Adapted Reading Passage (Base Context)
    if selected_sections.get("adapted_passage", True):
        log_status("Step 1/6: Generating adapted passage...")
        if source_type == "topic":
            prompt = f"""Role: Expert curriculum designer.
Task: Write an engaging, informative reading passage of about 250-350 words on the topic: '{source_val}' at a '{grade_level}' reading level in '{language}'.

Instructions:
Output a single JSON object containing only the key 'adapted_passage'. Do not include markdown code block wrappers.
{{
    "adapted_passage": "Your reading passage content..."
}}"""
        else:
            prompt = f"""Role: Expert curriculum designer.
Task: Adapt and rewrite the following text to match a '{grade_level}' reading level in '{language}' of about 250-350 words.
Source Text:
{source_val}

Instructions:
Output a single JSON object containing only the key 'adapted_passage'. Do not include markdown code block wrappers.
{{
    "adapted_passage": "Your rewritten reading passage..."
}}"""
        response, error = get_ai_response(prompt, ai_provider, ollama_model, api_key, ollama_url, gemini_model, format="json")
        data["adapted_passage"] = _clean_and_parse_json(response, "adapted_passage") or ""
    
    # Fallback if passage not generated or not requested: use original text (unless topic mode)
    if not data["adapted_passage"] and source_type != "topic":
        data["adapted_passage"] = source_val[:2000] # Cap context size

    if not data["adapted_passage"]:
        data["adapted_passage"] = f"An informational overview of the topic: {source_val}."

    # Step 2: Summary
    if selected_sections.get("summary", True):
        log_status("Step 2/6: Generating key summary...")
        prompt = f"""Role: Expert curriculum designer.
Task: Write 3 to 5 clear summary bullet points in '{language}' based on the reading passage.
Passage:
{data["adapted_passage"]}

Instructions:
Output a single JSON object containing only the key 'summary'. Do not include markdown code block wrappers.
{{
    "summary": [
        "First key point...",
        "Second key point...",
        "Third key point..."
    ]
}}"""
        response, error = get_ai_response(prompt, ai_provider, ollama_model, api_key, ollama_url, gemini_model, format="json")
        data["summary"] = _clean_and_parse_json(response, "summary") or []

    # Step 3: Vocabulary
    if selected_sections.get("vocab", True):
        log_status("Step 3/6: Extracting key vocabulary words...")
        prompt = f"""Role: Expert curriculum designer.
Task: Extract 4 key vocabulary words from the reading passage. For each term, provide its part of speech, a definition matching the grade level, and a context sentence in '{language}'.
Passage:
{data["adapted_passage"]}

Instructions:
Output a single JSON object containing only the key 'vocabulary'. Do not include markdown code block wrappers.
{{
    "vocabulary": [
        {{
            "word": "word",
            "pos": "noun/verb/adjective/adverb",
            "definition": "simple definition",
            "sentence": "context example sentence"
        }}
    ]
}}"""
        response, error = get_ai_response(prompt, ai_provider, ollama_model, api_key, ollama_url, gemini_model, format="json")
        data["vocabulary"] = _clean_and_parse_json(response, "vocabulary") or []

    # Step 4: MCQs
    if selected_sections.get("mcqs", True):
        log_status("Step 4/6: Generating multiple-choice questions...")
        prompt = f"""Role: Expert curriculum designer.
Task: Write 3 to 5 multiple-choice questions checking comprehension in '{language}' based on the reading passage. Include correct options and explanations.
Passage:
{data["adapted_passage"]}

Instructions:
Output a single JSON object containing only the key 'mcqs'. Do not include markdown code block wrappers.
{{
    "mcqs": [
        {{
            "question": "question text",
            "options": {{"A": "option A", "B": "option B", "C": "option C", "D": "option D"}},
            "correct_option": "A",
            "explanation": "why A is correct"
        }}
    ]
}}"""
        response, error = get_ai_response(prompt, ai_provider, ollama_model, api_key, ollama_url, gemini_model, format="json")
        data["mcqs"] = _clean_and_parse_json(response, "mcqs") or []

    # Step 5: Short Answer Questions
    if selected_sections.get("short_answers", True):
        log_status("Step 5/6: Generating short answer questions...")
        prompt = f"""Role: Expert curriculum designer.
Task: Write 3 short-answer comprehension questions and expected responses in '{language}' based on the reading passage.
Passage:
{data["adapted_passage"]}

Instructions:
Output a single JSON object containing only the key 'short_answers'. Do not include markdown code block wrappers.
{{
    "short_answers": [
        {{
            "question": "question text",
            "expected_answer": "model student answer"
        }}
    ]
}}"""
        response, error = get_ai_response(prompt, ai_provider, ollama_model, api_key, ollama_url, gemini_model, format="json")
        data["short_answers"] = _clean_and_parse_json(response, "short_answers") or []

    # Step 6: Bloom's Taxonomy Prompts
    if selected_sections.get("blooms", True):
        log_status("Step 6/6: Generating critical thinking prompts...")
        prompt = f"""Role: Expert curriculum designer.
Task: Write 3 open-ended discussion prompts in '{language}' based on the reading passage, matching Bloom's Taxonomy cognitive levels (e.g. Analyze, Evaluate, Create).
Passage:
{data["adapted_passage"]}

Instructions:
Output a single JSON object containing only the key 'blooms_prompts'. Do not include markdown code block wrappers.
{{
    "blooms_prompts": [
        {{
            "level": "Analyze/Evaluate/Create",
            "question": "critical open-ended question"
        }}
    ]
}}"""
        response, error = get_ai_response(prompt, ai_provider, ollama_model, api_key, ollama_url, gemini_model, format="json")
        data["blooms_prompts"] = _clean_and_parse_json(response, "blooms_prompts") or []

    log_status("Resource generation complete!")
    data["success"] = True
    return data

# =====================================================================
# WORD DOCUMENT (.DOCX) GENERATION FOR DIFFIT RESOURCES
# =====================================================================

def create_diffit_docx(data, title_topic, grade_level, language):
    """Generates the student-facing Exercise/Activity worksheet."""
    doc = docx.Document()
    
    # Page Setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Set default font
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Document Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title_p.add_run(f"Classroom Activity Worksheet: {title_topic}")
    t_run.font.size = Pt(18)
    t_run.bold = True
    
    # Document Subtitle
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run = sub_p.add_run(f"Target Reading Level: {grade_level}   |   Language: {language}\nName: _______________________   Date: _________________")
    s_run.font.size = Pt(10)
    s_run.italic = True
    doc.add_paragraph() # Spacer
    
    # 1. Adapted Reading Passage
    if data.get("adapted_passage"):
        h1 = doc.add_paragraph()
        r1 = h1.add_run("📖 Reading Passage")
        r1.font.size = Pt(14)
        r1.bold = True
        
        pass_p = doc.add_paragraph()
        pass_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pass_p.paragraph_format.line_spacing = 1.15
        pass_p.add_run(data["adapted_passage"])
        doc.add_paragraph()
        
    # 2. Key Summary (Bullet Points)
    if data.get("summary"):
        h2 = doc.add_paragraph()
        r2 = h2.add_run("📝 Article Summary")
        r2.font.size = Pt(14)
        r2.bold = True
        
        for bullet in data["summary"]:
            bp = doc.add_paragraph(style='List Bullet')
            bp.add_run(bullet)
        doc.add_paragraph()
        
    # 3. Vocabulary List (Formatted Table)
    if data.get("vocabulary"):
        h3 = doc.add_paragraph()
        r3 = h3.add_run("🔤 Vocabulary Study")
        r3.font.size = Pt(14)
        r3.bold = True
        
        # Add 4-column Table: Term, POS, Definition, Example Sentence
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Word / Term'
        hdr_cells[1].text = 'Part of Speech'
        hdr_cells[2].text = 'Definition'
        hdr_cells[3].text = 'Context Example'
        
        # Make headers bold
        for cell in hdr_cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
                    
        # Set column widths roughly
        widths = [Inches(1.2), Inches(1.0), Inches(2.3), Inches(2.0)]
        
        for item in data["vocabulary"]:
            row_cells = table.add_row().cells
            row_cells[0].text = item.get("word", "")
            row_cells[1].text = item.get("pos", "")
            row_cells[2].text = item.get("definition", "")
            row_cells[3].text = item.get("sentence", "")
            
        # Set widths
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
                
        doc.add_paragraph() # Spacer
        
    # 4. Comprehension Questions (Multiple Choice)
    if data.get("mcqs"):
        doc.add_page_break()
        h4 = doc.add_paragraph()
        r4 = h4.add_run("❓ Section A: Multiple-Choice Questions")
        r4.font.size = Pt(14)
        r4.bold = True
        
        for idx, mcq in enumerate(data["mcqs"], start=1):
            q_p = doc.add_paragraph()
            q_p.add_run(f"{idx}. ").bold = True
            q_p.add_run(mcq.get("question", ""))
            
            opts = mcq.get("options", {})
            for letter in ['A', 'B', 'C', 'D']:
                opt_val = opts.get(letter) or opts.get(letter.lower(), "")
                opt_p = doc.add_paragraph()
                opt_p.left_indent = Inches(0.4)
                opt_p.add_run(f"[{letter}] ").bold = True
                opt_p.add_run(opt_val)
            doc.add_paragraph()
            
    # 5. Short Answer Questions
    if data.get("short_answers"):
        h5 = doc.add_paragraph()
        r5 = h5.add_run("✍️ Section B: Short Answer Questions")
        r5.font.size = Pt(14)
        r5.bold = True
        
        for idx, sa in enumerate(data["short_answers"], start=1):
            q_p = doc.add_paragraph()
            q_p.add_run(f"{idx}. ").bold = True
            q_p.add_run(sa.get("question", ""))
            
            ans_p = doc.add_paragraph()
            ans_p.add_run("Answer: ___________________________________________________________________\n"
                          "___________________________________________________________________________").italic = True
            doc.add_paragraph()
            
    # 6. Bloom's Taxonomy Prompts
    if data.get("blooms_prompts"):
        h6 = doc.add_paragraph()
        r6 = h6.add_run("💡 Section C: Critical Thinking (Bloom's Taxonomy)")
        r6.font.size = Pt(14)
        r6.bold = True
        
        for idx, bp in enumerate(data["blooms_prompts"], start=1):
            q_p = doc.add_paragraph()
            q_p.add_run(f"{idx}. [{bp.get('level', 'Think')}]: ").bold = True
            q_p.add_run(bp.get("question", ""))
            
            ans_p = doc.add_paragraph()
            ans_p.add_run("Answer: ___________________________________________________________________\n"
                          "___________________________________________________________________________").italic = True
            doc.add_paragraph()

    # Footer
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.text = "Generated via CORTEX Quiz Generator"
    
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io


def create_diffit_answer_key_docx(data, title_topic, grade_level, language):
    """Generates the teacher-facing Answer Key document."""
    doc = docx.Document()
    
    # Page Setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title_p.add_run(f"TEACHER ANSWER KEY: {title_topic}")
    t_run.font.size = Pt(18)
    t_run.bold = True
    
    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run = sub_p.add_run(f"Target Reading Level: {grade_level}   |   Language: {language}")
    s_run.font.size = Pt(10)
    s_run.italic = True
    doc.add_paragraph()
    
    # MCQ Answers
    if data.get("mcqs"):
        h1 = doc.add_paragraph()
        r1 = h1.add_run("❓ Section A: Multiple-Choice Answer Key")
        r1.font.size = Pt(14)
        r1.bold = True
        
        for idx, mcq in enumerate(data["mcqs"], start=1):
            ans_p = doc.add_paragraph()
            ans_p.add_run(f"Question {idx}: ").bold = True
            ans_p.add_run(mcq.get("question", "") + "\n")
            
            correct = mcq.get("correct_option", "A").upper()
            opts = mcq.get("options", {})
            correct_val = opts.get(correct) or opts.get(correct.lower(), "")
            
            ans_p.add_run(f"👉 Correct Answer: [{correct}] {correct_val}\n").bold = True
            ans_p.add_run(f"Explanation: {mcq.get('explanation', 'None provided.')}").italic = True
            doc.add_paragraph()
            
    # Short Answer Answers
    if data.get("short_answers"):
        h2 = doc.add_paragraph()
        r2 = h2.add_run("✍️ Section B: Short Answer Expected Responses")
        r2.font.size = Pt(14)
        r2.bold = True
        
        for idx, sa in enumerate(data["short_answers"], start=1):
            ans_p = doc.add_paragraph()
            ans_p.add_run(f"Question {idx}: ").bold = True
            ans_p.add_run(sa.get("question", "") + "\n")
            ans_p.add_run(f"👉 Expected Answer: {sa.get('expected_answer', 'None provided.')}").bold = True
            doc.add_paragraph()
            
    # Bloom's taxonomy Reference
    if data.get("blooms_prompts"):
        h3 = doc.add_paragraph()
        r3 = h3.add_run("💡 Section C: Critical Thinking Questions (For Reference)")
        r3.font.size = Pt(14)
        r3.bold = True
        
        for idx, bp in enumerate(data["blooms_prompts"], start=1):
            ans_p = doc.add_paragraph()
            ans_p.add_run(f"Question {idx} [{bp.get('level', '')}]: ").bold = True
            ans_p.add_run(bp.get("question", ""))
            doc.add_paragraph()

    # Footer
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.text = "Generated via CORTEX Quiz Generator"
    
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io


def generate_grammar_exercises(
    passage,
    ai_provider=None,
    ollama_model=None,
    api_key=None,
    ollama_url=None,
    gemini_model=None
):
    """
    Algorithmic generation of grammar/spelling exercises containing three parts:
    1. Part 1: Spelling & Word Completion (Gapped word cloze) - 5 items.
    2. Part 2: True / False same-POS Vocabulary Exercises - 5 items.
    3. Part 3: Multiple-Choice Questions (MCQ) with same-POS distractors - 5 items.
    """
    import re
    import random
    import os
    import glob
    import tempfile
    import duckdb

    if not passage:
        return None

    # Clean up whitespace and chunk into sentences
    text = re.sub(r'\s+', ' ', passage).strip()
    sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+', text) if len(s.strip()) > 15]
    
    if not sentences:
        return None

    # Shuffle sentences
    random.shuffle(sentences)
    
    # ----------------------------------------------------
    # PART 1: GAPPED WORDS (CLOZE) - 5 ITEMS
    # ----------------------------------------------------
    fill_in_blanks = []
    part1_sents = []
    
    for sent in sentences:
        if len(fill_in_blanks) >= 5:
            break
            
        words = re.findall(r'\b[a-zA-Z]{6,}\b', sent)
        if not words:
            continue
            
        target_word = random.choice(words)
        n_letters = len(target_word)
        remove_count = int(round(n_letters * random.uniform(0.5, 0.6)))
        remove_count = max(1, min(remove_count, n_letters - 2))
        
        indices_to_remove = random.sample(range(n_letters), remove_count)
        
        gapped_chars = list(target_word)
        for idx in indices_to_remove:
            gapped_chars[idx] = ' '
        gapped_word = "".join(gapped_chars)
        
        pattern = r'\b' + re.escape(target_word) + r'\b'
        sentence_with_gap = re.sub(pattern, f"`{gapped_word}`", sent, count=1)
        
        fill_in_blanks.append({
            "sentence": sentence_with_gap,
            "base_word": gapped_word,
            "correct_answer": target_word,
            "explanation": f"Spelling and vocabulary practice for the word '{target_word}' from the reading passage."
        })
        part1_sents.append(sent)

    # ----------------------------------------------------
    # PART 2 & 3 DATABASE POS RETRIEVAL
    # ----------------------------------------------------
    # Try to resolve active DB path from streamlit state
    db_path = None
    try:
        from ui_streamlit.state_manager import get_state
        db_path = get_state('current_corpus_path')
    except Exception:
        pass

    # Fallback to temp dir
    if not db_path or not os.path.exists(db_path):
        temp_dir = tempfile.gettempdir()
        db_files = glob.glob(os.path.join(temp_dir, "corpus_*.duckdb"))
        if db_files:
            db_path = max(db_files, key=os.path.getmtime)

    # Load POS mapping from DuckDB
    pos_map = {} # POS -> list of tokens
    word_to_pos = {} # token_low -> POS
    
    if db_path and os.path.exists(db_path):
        try:
            con = duckdb.connect(db_path, read_only=True)
            # Fetch distinct tokens and their POS tags
            rows = con.execute("""
                SELECT DISTINCT token, pos 
                FROM corpus 
                WHERE pos IS NOT NULL 
                  AND LENGTH(token) > 4
                  AND regexp_matches(token, '^[a-zA-Z]+$')
            """).fetchall()
            con.close()
            
            for token, pos in rows:
                token_clean = token.strip()
                pos_clean = pos.strip().upper()
                if not token_clean or not pos_clean:
                    continue
                word_to_pos[token_clean.lower()] = pos_clean
                if pos_clean not in pos_map:
                    pos_map[pos_clean] = []
                pos_map[pos_clean].append(token_clean)
        except Exception as e:
            print(f"[ERROR] Database POS retrieval failed: {e}")

    # Fallback POS map in case DB has no POS or is missing
    if not pos_map:
        pos_map = {
            "NOUN": ["narrative", "experience", "observation", "conversation", "reflection", "routine", "purpose", "weather", "atmosphere", "responsibility", "community"],
            "VERB": ["describes", "moved", "made", "changed", "developed", "responded", "focused", "spent", "observing", "discussing", "created", "exchanged"],
            "ADJ": ["ordinary", "natural", "unexpected", "memorable", "gradual", "practical", "realistic", "broader", "personal", "coherent", "standard"],
            "ADV": ["gradually", "naturally", "normally", "clearly", "practically", "especially", "unexpectedly"]
        }
        for pos, words in pos_map.items():
            for w in words:
                word_to_pos[w.lower()] = pos

    # ----------------------------------------------------
    # PART 2: TRUE / FALSE (SAME POS) - 5 ITEMS
    # ----------------------------------------------------
    true_false = []
    part2_sents = []
    remaining_for_tf = [s for s in sentences if s not in part1_sents]
    if len(remaining_for_tf) < 5:
        remaining_for_tf = sentences
        
    for sent in remaining_for_tf:
        if len(true_false) >= 5:
            break
            
        # Find all words in this sentence that have known POS and length > 4
        words_in_sent = re.findall(r'\b[a-zA-Z]{5,}\b', sent)
        candidate_words = [w for w in words_in_sent if w.lower() in word_to_pos]
        
        if not candidate_words:
            continue
            
        target_word = random.choice(candidate_words)
        target_pos = word_to_pos[target_word.lower()]
        
        # Get alternative distractors with same POS
        distractions = [w for w in pos_map.get(target_pos, []) if w.lower() != target_word.lower()]
        
        # Determine if this item will be True (correct word) or False (incorrect word)
        is_correct = random.choice([True, False]) if distractions else True
        
        if is_correct:
            displayed_word = target_word
            explanation = f"Correct. The highlighted word '**{target_word}**' matches the original text from the reading passage."
        else:
            distract_word = random.choice(distractions)
            # Match casing of original word
            if target_word[0].isupper():
                distract_word = distract_word.capitalize()
            else:
                distract_word = distract_word.lower()
            displayed_word = distract_word
            explanation = f"Incorrect. The highlighted word '**{displayed_word}**' is wrong. The correct word from the reading passage was '**{target_word}**'."

        # Replace first occurrence of target word with highlighted displayed word (using high-contrast HTML span tag)
        pattern = r'\b' + re.escape(target_word) + r'\b'
        highlighted_html = f'<span style="background-color: rgba(0, 173, 181, 0.18); color: #00FFF5; padding: 2px 6px; border-radius: 4px; font-weight: bold; border: 1px solid rgba(0, 173, 181, 0.5); font-size: 1.05rem;">{displayed_word}</span>'
        sentence_highlighted = re.sub(pattern, highlighted_html, sent, count=1)
        
        true_false.append({
            "sentence": sentence_highlighted,
            "original_word": target_word,
            "displayed_word": displayed_word,
            "is_correct": is_correct,
            "explanation": explanation
        })
        part2_sents.append(sent)

    # ----------------------------------------------------
    # PART 3: MULTIPLE-CHOICE QUESTIONS (MCQ) - 5 ITEMS
    # ----------------------------------------------------
    mcqs = []
    remaining_for_mcq = [s for s in sentences if s not in part1_sents and s not in part2_sents]
    if len(remaining_for_mcq) < 5:
        remaining_for_mcq = sentences

    for sent in remaining_for_mcq:
        if len(mcqs) >= 5:
            break
            
        words_in_sent = re.findall(r'\b[a-zA-Z]{5,}\b', sent)
        candidate_words = [w for w in words_in_sent if w.lower() in word_to_pos]
        
        if not candidate_words:
            continue
            
        target_word = random.choice(candidate_words)
        target_pos = word_to_pos[target_word.lower()]
        
        # Get alternative distractors with same POS
        distractions = [w for w in pos_map.get(target_pos, []) if w.lower() != target_word.lower()]
        
        # We need exactly 4 distractors (to make 5 options in total)
        # If not enough, pad with words from other POS groups
        if len(distractions) < 4:
            all_other_words = []
            for other_pos, other_words in pos_map.items():
                all_other_words.extend([w for w in other_words if w.lower() != target_word.lower()])
            random.shuffle(all_other_words)
            distractions.extend(all_other_words)
            distractions = list(dict.fromkeys(distractions)) # Dedup
            
        selected_distractions = random.sample(distractions, min(4, len(distractions)))
        while len(selected_distractions) < 4:
            selected_distractions.append("word")
            
        # Match case
        cased_options = []
        is_upper = target_word[0].isupper()
        for w in selected_distractions:
            cased_options.append(w.capitalize() if is_upper else w.lower())
            
        # Combine target word and distractors
        cased_options.append(target_word)
        random.shuffle(cased_options)
        
        # Create letters options
        letters = ['A', 'B', 'C', 'D', 'E']
        options_dict = {}
        correct_letter = 'A'
        for idx, opt in enumerate(cased_options[:5]):
            letter = letters[idx]
            options_dict[letter] = opt
            if opt.lower() == target_word.lower():
                correct_letter = letter
                
        # Completely remove target word and replace with a dashed cyan blank line
        pattern = r'\b' + re.escape(target_word) + r'\b'
        blank_html = '<span style="border-bottom: 2px dashed #00FFF5; padding-left: 30px; padding-right: 30px; font-weight: bold; color: #00FFF5;">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span>'
        sentence_with_blank = re.sub(pattern, blank_html, sent, count=1)
        
        mcqs.append({
            "question": sentence_with_blank,
            "options": options_dict,
            "correct_option": correct_letter,
            "explanation": f"The correct word to fill the blank is '**{target_word}**' (POS: {target_pos}) based on the reading passage."
        })
        
    return {
        "fill_in_blanks": fill_in_blanks,
        "true_false": true_false,
        "mcqs": mcqs
    }

def _modify_number(sentence):
    import re
    import random
    # Try digits first
    digit_matches = list(re.finditer(r'\b\d+\b', sentence))
    if digit_matches:
        match = random.choice(digit_matches)
        num_str = match.group()
        try:
            num_val = int(num_str)
            if num_val > 1000:
                new_val = num_val + random.choice([-10, -5, 5, 10, 100])
            else:
                new_val = num_val + random.choice([-2, -1, 1, 2, 3])
                if new_val < 0:
                    new_val = num_val + 2
            start, end = match.span()
            return sentence[:start] + str(new_val) + sentence[end:]
        except ValueError:
            pass
            
    # Try word numbers
    num_words = {
        "one": "three", "two": "five", "three": "six", "four": "eight", "five": "nine",
        "six": "ten", "seven": "eleven", "eight": "twelve", "nine": "fifteen", "ten": "twenty",
        "eleven": "thirteen", "twelve": "fourteen", "twenty": "thirty", "thirty": "fifty",
        "forty": "sixty", "fifty": "seventy", "hundred": "thousand", "thousand": "million",
        "first": "third", "second": "fourth", "third": "fifth", "fourth": "sixth", "fifth": "seventh"
    }
    for w, rep in num_words.items():
        pattern = r'\b' + re.escape(w) + r'\b'
        match = re.search(pattern, sentence, re.IGNORECASE)
        if match:
            orig_word = match.group()
            new_word = rep
            if orig_word[0].isupper():
                new_word = rep.capitalize()
            return re.sub(pattern, new_word, sentence, count=1, flags=re.IGNORECASE)
            
    return None

def _change_negative_to_positive(sentence):
    import re
    neg_map = {
        r"\bnever\b": "always",
        r"\bno\b": "some",
        r"\bnothing\b": "something",
        r"\bnobody\b": "somebody",
        r"\bnowhere\b": "somewhere",
        r"\bnone\b": "all",
        r"\bneither\b": "either",
        r"\bnor\b": "or",
        r"\bbarely\b": "fully",
        r"\bscarcely\b": "fully",
        r"\bhardly\b": "completely",
        r"\bdon't\b": "do",
        r"\bdoesn't\b": "does",
        r"\bdidn't\b": "did",
        r"\bwon't\b": "will",
        r"\bcan't\b": "can",
        r"\bcannot\b": "can",
        r"\bshouldn't\b": "should",
        r"\bwouldn't\b": "would",
        r"\bisn't\b": "is",
        r"\baren't\b": "are",
        r"\bwasn't\b": "was",
        r"\bweren't\b": "were",
        r"\bhasn't\b": "has",
        r"\bhaven't\b": "have",
        r"\bhadn't\b": "had"
    }
    for pat, rep in neg_map.items():
        if re.search(pat, sentence, re.IGNORECASE):
            match = re.search(pat, sentence, re.IGNORECASE)
            orig_word = match.group()
            new_word = rep
            if orig_word[0].isupper():
                new_word = rep.capitalize()
            return re.sub(pat, new_word, sentence, count=1, flags=re.IGNORECASE)
            
    # Try general "not" removal
    if re.search(r'\bnot\b', sentence, re.IGNORECASE):
        sentence_clean = re.sub(r'\bdo not\b', 'do', sentence, count=1, flags=re.IGNORECASE)
        sentence_clean = re.sub(r'\bdoes not\b', 'does', sentence_clean, count=1, flags=re.IGNORECASE)
        sentence_clean = re.sub(r'\bdid not\b', 'did', sentence_clean, count=1, flags=re.IGNORECASE)
        if sentence_clean != sentence:
            return sentence_clean
        # Fallback simple remove not
        sentence_clean = re.sub(r'\s+not\b', '', sentence, count=1, flags=re.IGNORECASE)
        sentence_clean = re.sub(r'\bnot\s+', '', sentence_clean, count=1, flags=re.IGNORECASE)
        return sentence_clean
        
    return None

def _change_degree_adverb(sentence, word_to_pos, pos_map):
    import re
    import random
    adv_pairs = {
        "always": "never",
        "never": "always",
        "extremely": "slightly",
        "highly": "slightly",
        "very": "somewhat",
        "completely": "partially",
        "absolutely": "partially",
        "totally": "partially",
        "significantly": "marginally",
        "greatly": "marginally",
        "frequently": "rarely",
        "often": "seldom",
        "rarely": "frequently",
        "seldom": "often"
    }
    for w, rep in adv_pairs.items():
        pattern = r'\b' + re.escape(w) + r'\b'
        if re.search(pattern, sentence, re.IGNORECASE):
            match = re.search(pattern, sentence, re.IGNORECASE)
            orig_word = match.group()
            new_word = rep
            if orig_word[0].isupper():
                new_word = rep.capitalize()
            return re.sub(pattern, new_word, sentence, count=1, flags=re.IGNORECASE)
            
    # Fallback to general adverb replacement
    words = re.findall(r'\b[a-zA-Z]+\b', sentence)
    adv_words = [w for w in words if w.lower() in word_to_pos and word_to_pos[w.lower()] == 'ADV']
    if adv_words:
        target_adv = random.choice(adv_words)
        alternatives = [w for w in pos_map.get('ADV', []) if w.lower() != target_adv.lower()]
        if not alternatives:
            alternatives = ["silently", "quickly", "gradually", "naturally", "clearly", "unexpectedly"]
        rep = random.choice(alternatives)
        if target_adv[0].isupper():
            rep = rep.capitalize()
        else:
            rep = rep.lower()
        pattern = r'\b' + re.escape(target_adv) + r'\b'
        return re.sub(pattern, rep, sentence, count=1)
        
    return None

def _swap_proper_noun(sentence, all_proper_nouns, word_to_pos):
    import re
    import random
    words_in_sent = re.findall(r'\b[a-zA-Z]+\b', sentence)
    candidate_props = []
    
    for idx, w in enumerate(words_in_sent):
        is_prop = False
        if w.lower() in word_to_pos and word_to_pos[w.lower()] == 'PROPN':
            is_prop = True
        elif idx > 0 and w[0].isupper():
            is_prop = True
            
        if is_prop:
            candidate_props.append(w)
            
    if candidate_props:
        target_prop = random.choice(candidate_props)
        alternatives = [p for p in all_proper_nouns if p.lower() != target_prop.lower()]
        if alternatives:
            rep = random.choice(alternatives)
            pattern = r'\b' + re.escape(target_prop) + r'\b'
            return re.sub(pattern, rep, sentence, count=1)
            
    return None

def _fallback_modify_sentence(sentence, word_to_pos, pos_map):
    import re
    import random
    words = re.findall(r'\b[a-zA-Z]+\b', sentence)
    
    # Try noun replacement
    noun_words = [w for w in words if w.lower() in word_to_pos and word_to_pos[w.lower()] == 'NOUN']
    if noun_words:
        target_noun = random.choice(noun_words)
        alternatives = [w for w in pos_map.get('NOUN', []) if w.lower() != target_noun.lower()]
        if not alternatives:
            alternatives = ["object", "thing", "event", "situation", "concept", "element"]
        rep = random.choice(alternatives)
        if target_noun[0].isupper():
            rep = rep.capitalize()
        else:
            rep = rep.lower()
        pattern = r'\b' + re.escape(target_noun) + r'\b'
        return re.sub(pattern, rep, sentence, count=1)
        
    # Try adjective replacement
    adj_words = [w for w in words if w.lower() in word_to_pos and word_to_pos[w.lower()] == 'ADJ']
    if adj_words:
        target_adj = random.choice(adj_words)
        alternatives = [w for w in pos_map.get('ADJ', []) if w.lower() != target_adj.lower()]
        if not alternatives:
            alternatives = ["different", "similar", "important", "major", "minor", "significant"]
        rep = random.choice(alternatives)
        if target_adj[0].isupper():
            rep = rep.capitalize()
        else:
            rep = rep.lower()
        pattern = r'\b' + re.escape(target_adj) + r'\b'
        return re.sub(pattern, rep, sentence, count=1)
        
    if words:
        long_words = [w for w in words if len(w) > 3]
        if long_words:
            target = random.choice(long_words)
            pattern = r'\b' + re.escape(target) + r'\b'
            return re.sub(pattern, '', sentence, count=1).replace('  ', ' ').strip()
            
    return sentence + " (not)"

def generate_reading_comprehension(
    passage,
    ai_provider=None,
    ollama_model=None,
    api_key=None,
    ollama_url=None,
    gemini_model=None
):
    """
    Algorithmic generation of Reading Comprehension exercises:
    1. Type 1: 'Which of the following is mentioned in the text?'
       Correct option: Exact unmodified sentence.
       4 Distractors: number modification, negative-to-positive, degree adverb modification, proper noun swap.
    2. Type 2: 'Which of the following is NOT mentioned in the text?'
       Correct option: A modified distractor sentence from Type 1.
       4 Distractors: Unmodified exact sentences from the text.
    3. Type 3: Pronoun Resolution (Coreference).
       Identifies a pronoun (they, it, he, she, we) and resolves its reference.
    4. Type 4: Paragraph Start/End Completion.
       Selects 5 consecutive sentences. Displays 4, asks for the 1st (start) or 5th (end) sentence.
    """
    import re
    import random
    import os
    import glob
    import tempfile
    import duckdb
    import spacy

    if not passage:
        return None

    # Clean up whitespace and chunk into sentences (keeping original sequential order)
    text = re.sub(r'\s+', ' ', passage).strip()
    original_sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+', text) if len(s.strip()) > 15]
    
    num_sentences = len(original_sentences)
    type_4_skipped = num_sentences < 6

    # For safety, let's ensure we have a fallback pool of sentences
    sentences = original_sentences[:]
    if len(sentences) < 5:
        sentences = sentences * 3

    # Load POS database
    db_path = None
    try:
        from ui_streamlit.state_manager import get_state
        db_path = get_state('current_corpus_path')
    except Exception:
        pass

    if not db_path or not os.path.exists(db_path):
        temp_dir = tempfile.gettempdir()
        db_files = glob.glob(os.path.join(temp_dir, "corpus_*.duckdb"))
        if db_files:
            db_path = max(db_files, key=os.path.getmtime)

    # Database Schema Check/Migration:
    if db_path and os.path.exists(db_path):
        try:
            con_rw = duckdb.connect(db_path, read_only=False)
            columns = con_rw.execute("PRAGMA table_info(corpus)").fetchall()
            col_names = [col[1] for col in columns]
            if "ent_type" not in col_names:
                con_rw.execute("ALTER TABLE corpus ADD COLUMN ent_type VARCHAR DEFAULT ''")
            con_rw.close()
        except Exception as e:
            print(f"[WARNING] Could not execute DB migration: {e}")

    pos_map = {}
    word_to_pos = {}
    
    if db_path and os.path.exists(db_path):
        try:
            con = duckdb.connect(db_path, read_only=True)
            rows = con.execute("""
                SELECT DISTINCT token, pos 
                FROM corpus 
                WHERE pos IS NOT NULL 
                  AND LENGTH(token) > 2
                  AND regexp_matches(token, '^[a-zA-Z]+$')
            """).fetchall()
            con.close()
            
            for token, pos in rows:
                token_clean = token.strip()
                pos_clean = pos.strip().upper()
                if not token_clean or not pos_clean:
                    continue
                word_to_pos[token_clean.lower()] = pos_clean
                if pos_clean not in pos_map:
                    pos_map[pos_clean] = []
                pos_map[pos_clean].append(token_clean)
        except Exception as e:
            print(f"[ERROR] Database POS retrieval failed: {e}")

    # Fallbacks for POS lists
    if not pos_map:
        pos_map = {
            "NOUN": ["narrative", "experience", "observation", "conversation", "reflection", "routine", "purpose", "weather", "atmosphere", "responsibility", "community"],
            "VERB": ["describes", "moved", "made", "changed", "developed", "responded", "focused", "spent", "observing", "discussing", "created", "exchanged"],
            "ADJ": ["ordinary", "natural", "unexpected", "memorable", "gradual", "practical", "realistic", "broader", "personal", "coherent", "standard"],
            "ADV": ["gradually", "naturally", "normally", "clearly", "practically", "especially", "unexpectedly"]
        }
        for pos, words in pos_map.items():
            for w in words:
                word_to_pos[w.lower()] = pos

    # Gather proper nouns in the passage
    all_proper_nouns = set()
    for sent in sentences:
        words = re.findall(r'\b[a-zA-Z]+\b', sent)
        for idx, w in enumerate(words):
            if w.lower() in word_to_pos and word_to_pos[w.lower()] == 'PROPN':
                all_proper_nouns.add(w)
            elif idx > 0 and w[0].isupper():
                common_pronouns_starts = {"They", "We", "The", "He", "She", "It", "You", "I", "This", "That", "These", "Those", "There", "Here"}
                if w not in common_pronouns_starts:
                    all_proper_nouns.add(w)

    # Load SpaCy pipeline once
    nlp = None
    doc = None
    try:
        nlp = spacy.load("en_core_web_sm")
        doc = nlp(passage)
    except Exception as e:
        print(f"[WARNING] SpaCy load failed: {e}")

    questions = []

    # ----------------------------------------------------
    # RUN PRONOUN RESOLUTION (Type 3)
    # ----------------------------------------------------
    type_3_questions = []
    try:
        
        # 1. Extract conjoined noun phrases
        noun_phrases = []
        visited = set()
        for chunk in doc.noun_chunks:
            root = chunk.root
            if root in visited:
                continue
            if root.dep_ == "conj":
                continue
            conjuncts = [chunk]
            visited.add(root)
            
            def add_conjuncts(tok):
                for child in tok.children:
                    if child.dep_ == "conj":
                        child_chunk = next((c for c in doc.noun_chunks if c.root == child), None)
                        if child_chunk:
                            conjuncts.append(child_chunk)
                            visited.add(child)
                            add_conjuncts(child)
            add_conjuncts(root)
            
            if len(conjuncts) > 1:
                start_char = conjuncts[0].start_char
                end_char = conjuncts[-1].end_char
                noun_phrases.append({
                    "text": doc.text[start_char:end_char],
                    "is_plural": True,
                    "is_proper": any(c.root.pos_ == "PROPN" for c in conjuncts),
                    "last_token_idx": conjuncts[-1].root.i,
                    "root_dep": root.dep_,
                    "sent_start_char": chunk.sent.start_char,
                    "sent_end_char": chunk.sent.end_char
                })
            else:
                is_plural = (chunk.root.tag_ in ["NNS", "NNPS"]) or (chunk.text.lower() in ["we", "they", "them", "us", "people", "children", "men", "women"])
                noun_phrases.append({
                    "text": chunk.text,
                    "is_plural": is_plural,
                    "is_proper": chunk.root.pos_ == "PROPN",
                    "last_token_idx": root.i,
                    "root_dep": root.dep_,
                    "sent_start_char": chunk.sent.start_char,
                    "sent_end_char": chunk.sent.end_char
                })

        # 2. Score and resolve pronouns
        resolutions = []
        for token in doc:
            pron_lower = token.text.lower()
            if token.pos_ != "PRON" or pron_lower not in ["they", "them", "their", "it", "its", "he", "him", "his", "she", "her", "we"]:
                continue
            is_plural_pron = pron_lower in ["they", "them", "their", "we"]
            
            candidates = []
            for np in noun_phrases:
                if np["last_token_idx"] >= token.i:
                    continue
                pron_sent = token.sent
                is_same_sent = (np["sent_start_char"] == pron_sent.start_char)
                prev_sent = None
                sents = list(doc.sents)
                for s_idx, s in enumerate(sents):
                    if s.start_char == pron_sent.start_char:
                        if s_idx > 0:
                            prev_sent = sents[s_idx - 1]
                        break
                is_prev_sent = (prev_sent and np["sent_start_char"] == prev_sent.start_char)
                if not (is_same_sent or is_prev_sent):
                    continue
                if is_plural_pron != np["is_plural"]:
                    continue
                score = 0
                if is_same_sent: score += 3
                elif is_prev_sent: score += 1
                if np["root_dep"] in ["nsubj", "nsubjpass"]: score += 5
                elif np["root_dep"] == "dobj": score += 3
                elif np["root_dep"] == "pobj": score += 2
                dist = token.i - np["last_token_idx"]
                score += max(0, 10 - (dist // 3))
                candidates.append((score, np))
                
            if candidates:
                candidates.sort(key=lambda x: x[0], reverse=True)
                best_score, best_np = candidates[0]
                resolutions.append((token, best_np, noun_phrases))

        # 3. Create Type 3 Questions
        for token, np_correct, all_nps in resolutions:
            correct_text = np_correct["text"]
            pronoun = token.text
            sentence = token.sent.text.strip()
            
            # Formulate prompt
            q_text = f'In the sentence: "{sentence}", the pronoun "{pronoun}" refers to...'
            
            # Select distractors
            candidates = [np for np in all_nps if np["text"].lower() != correct_text.lower() and np["text"].lower() != pronoun.lower() and not np["text"].lower() in ["we", "they", "them", "us", "he", "him", "she", "her", "it", "its", "i", "me", "you", "their", "his", "hers"]]
            
            distractors = []
            
            if np_correct["is_proper"]:
                props = [c for c in candidates if c["is_proper"]]
                props.sort(key=lambda x: abs(token.i - x["last_token_idx"]))
                for p in props:
                    if p["text"] not in distractors:
                        distractors.append(p["text"])
                        
            matching_num = [c for c in candidates if c["is_plural"] == np_correct["is_plural"]]
            matching_num.sort(key=lambda x: abs(token.i - x["last_token_idx"]))
            for m in matching_num:
                if len(distractors) >= 4:
                    break
                if m["text"] not in distractors:
                    distractors.append(m["text"])
                    
            candidates.sort(key=lambda x: abs(token.i - x["last_token_idx"]))
            for c in candidates:
                if len(distractors) >= 4:
                    break
                if c["text"] not in distractors:
                    distractors.append(c["text"])
                    
            # Fallback padding
            synthetic_props = ["John", "Mary", "Paris", "London", "the supervisor", "the company"]
            synthetic_nouns = ["the system", "the results", "the process", "the documents", "the boxes"]
            fallback_pool = synthetic_props if np_correct["is_proper"] else synthetic_nouns
            for f in fallback_pool:
                if len(distractors) >= 4:
                    break
                if f.lower() != correct_text.lower() and f not in distractors:
                    distractors.append(f)
                    
            distractors = distractors[:4]
            all_opts = [correct_text] + distractors
            random.shuffle(all_opts)
            
            letters = ['A', 'B', 'C', 'D', 'E']
            options_dict = {}
            correct_letter = 'A'
            for idx, opt in enumerate(all_opts[:5]):
                letter = letters[idx]
                options_dict[letter] = opt
                if opt == correct_text:
                    correct_letter = letter
                    
            type_3_questions.append({
                "type": "type_3",
                "question": q_text,
                "options": options_dict,
                "correct_option": correct_letter,
                "correct_text": correct_text,
                "explanation": f"The pronoun '{pronoun}' refers back to the antecedent '{correct_text}'."
            })
    except Exception as e:
        print(f"[ERROR] SpaCy Pronoun Resolution failed: {e}")

    # Shuffled Type 3 questions
    random.shuffle(type_3_questions)

    # ----------------------------------------------------
    # Helper for Type 1 & 2 distractor builders
    # ----------------------------------------------------
    def get_distractor_sentence(sent_pool, dist_type):
        for s in sent_pool:
            res_s = None
            if dist_type == "number":
                res_s = _modify_number(s)
            elif dist_type == "negative":
                res_s = _change_negative_to_positive(s)
            elif dist_type == "adverb":
                res_s = _change_degree_adverb(s, word_to_pos, pos_map)
            elif dist_type == "proper_noun":
                res_s = _swap_proper_noun(s, all_proper_nouns, word_to_pos)
                
            if res_s and res_s != s:
                return s, res_s, dist_type
        return None

    def build_distractor_with_fallback(sent_pool, preferred_type, successful_types):
        source_pool = sent_pool if sent_pool else sentences
        res = get_distractor_sentence(source_pool, preferred_type)
        if res:
            return res
        for other_type in successful_types:
            res = get_distractor_sentence(source_pool, other_type)
            if res:
                return res
        for s in source_pool:
            res_s = _fallback_modify_sentence(s, word_to_pos, pos_map)
            if res_s and res_s != s:
                return s, res_s, "fallback"
        return source_pool[0], source_pool[0] + " (modified)", "fallback"

    # ----------------------------------------------------
    # RUN PARAGRAPH COMPLETION (Type 4)
    # ----------------------------------------------------
    type_4_questions = []
    if not type_4_skipped:
        try:
            # We generate 1 start question and 1 end question
            # Select 5 consecutive sentences
            start_idx = random.randint(0, num_sentences - 5)
            segment_sents = original_sentences[start_idx:start_idx + 5]
            remaining_sents = [s for s in original_sentences if s not in segment_sents]
            
            # Start Completion Question
            s_correct = segment_sents[0]
            displayed_start = " ".join(segment_sents[1:])
            q_start_prompt = (
                "Which of the sentences below best fits the start of the segment?\n\n"
                f"> {displayed_start}"
            )
            
            d_start = _generate_paragraph_distractors(remaining_sents, segment_sents, 4, all_proper_nouns, word_to_pos, pos_map)
            opts_start = [s_correct] + d_start
            random.shuffle(opts_start)
            
            letters = ['A', 'B', 'C', 'D', 'E']
            options_dict_start = {}
            correct_letter_start = 'A'
            for idx, opt in enumerate(opts_start[:5]):
                letter = letters[idx]
                options_dict_start[letter] = opt
                if opt == s_correct:
                    correct_letter_start = letter
                    
            type_4_questions.append({
                "type": "type_4_start",
                "question": q_start_prompt,
                "options": options_dict_start,
                "correct_option": correct_letter_start,
                "correct_text": s_correct,
                "explanation": f"The segment starts with: '{s_correct}'."
            })
            
            # End Completion Question
            if num_sentences >= 10:
                end_start_idx = random.randint(0, num_sentences - 5)
                if end_start_idx == start_idx:
                    end_start_idx = (start_idx + 2) % (num_sentences - 4)
                segment_sents_end = original_sentences[end_start_idx:end_start_idx + 5]
            else:
                segment_sents_end = segment_sents
                
            remaining_sents_end = [s for s in original_sentences if s not in segment_sents_end]
            
            s_correct_end = segment_sents_end[-1]
            displayed_end = " ".join(segment_sents_end[:-1])
            q_end_prompt = (
                "Which of the sentences below best fits the end of the segment?\n\n"
                f"> {displayed_end}"
            )
            
            d_end = _generate_paragraph_distractors(remaining_sents_end, segment_sents_end, 4, all_proper_nouns, word_to_pos, pos_map)
            opts_end = [s_correct_end] + d_end
            random.shuffle(opts_end)
            
            options_dict_end = {}
            correct_letter_end = 'A'
            for idx, opt in enumerate(opts_end[:5]):
                letter = letters[idx]
                options_dict_end[letter] = opt
                if opt == s_correct_end:
                    correct_letter_end = letter
                    
            type_4_questions.append({
                "type": "type_4_end",
                "question": q_end_prompt,
                "options": options_dict_end,
                "correct_option": correct_letter_end,
                "correct_text": s_correct_end,
                "explanation": f"The segment ends with: '{s_correct_end}'."
            })
        except Exception as e:
            print(f"[ERROR] Type 4 question generation failed: {e}")
            type_4_skipped = True

    # ----------------------------------------------------
    # RUN NAMED ENTITY RECOGNITION (Type 5)
    # ----------------------------------------------------
    type_5_questions = []
    if doc:
        try:
            NER_MAP = {
                "PERSON": ("Who", "person"),
                "GPE": ("Where", "location"),
                "LOC": ("Where", "location"),
                "FAC": ("Where", "location"),
                "DATE": ("When", "time"),
                "TIME": ("When", "time"),
                "ORG": ("What", "organization"),
                "PRODUCT": ("What", "object"),
                "EVENT": ("What", "event"),
                "WORK_OF_ART": ("What", "work_of_art")
            }
            
            valid_entities = []
            for ent in doc.ents:
                label = ent.label_
                if label in NER_MAP:
                    ent_text = ent.text.strip()
                    if len(ent_text) > 1 and not ent_text.lower() in ["we", "they", "them", "us", "he", "him", "she", "her", "it", "its", "i", "me", "you", "their", "his", "hers"]:
                        valid_entities.append(ent)
            
            for ent in valid_entities:
                correct_text = ent.text.strip()
                label = ent.label_
                q_word, category = NER_MAP[label]
                
                sent_obj = ent.sent
                sent_start_char = sent_obj.start_char
                ent_start_in_sent = ent.start_char - sent_start_char
                ent_end_in_sent = ent.end_char - sent_start_char
                
                if ent_start_in_sent < 0 or ent_end_in_sent > len(sent_obj.text):
                    continue
                    
                sentence_with_blank = sent_obj.text[:ent_start_in_sent] + "[Blank]" + sent_obj.text[ent_end_in_sent:]
                sentence_with_blank = re.sub(r'\s+', ' ', sentence_with_blank).strip()
                
                q_prompt = f"According to the passage, {q_word.lower()} or what {category} completes the statement below?\n\n\"{sentence_with_blank}\""
                
                distractors = []
                
                # 1. Current passage same entity type
                same_type_passage = [e.text.strip() for e in valid_entities if e.label_ == label and e.text.strip().lower() != correct_text.lower()]
                for d in same_type_passage:
                    if d not in distractors and len(distractors) < 4:
                        distractors.append(d)
                
                # 2. Database same entity type query
                if len(distractors) < 4 and db_path and os.path.exists(db_path):
                    try:
                        con = duckdb.connect(db_path, read_only=True)
                        columns = con.execute("PRAGMA table_info(corpus)").fetchall()
                        col_names = [col[1] for col in columns]
                        if "ent_type" in col_names:
                            db_ents = con.execute("""
                                SELECT DISTINCT token 
                                FROM corpus 
                                WHERE ent_type = ? 
                                  AND token IS NOT NULL 
                                  AND LENGTH(token) > 1
                                  AND token != ?
                                LIMIT 20
                            """, [label, correct_text]).fetchall()
                            for row in db_ents:
                                d = row[0].strip()
                                if d not in distractors and len(distractors) < 4:
                                    distractors.append(d)
                        con.close()
                    except:
                        pass
                
                # 3. Database POS fallback query
                if len(distractors) < 4 and db_path and os.path.exists(db_path):
                    try:
                        con = duckdb.connect(db_path, read_only=True)
                        if q_word in ["Who", "Where"] or label == "ORG":
                            db_prop = con.execute("""
                                SELECT DISTINCT token 
                                FROM corpus 
                                WHERE pos = 'PROPN' 
                                  AND token IS NOT NULL 
                                  AND LENGTH(token) > 2
                                  AND token != ?
                                LIMIT 20
                            """, [correct_text]).fetchall()
                            for row in db_prop:
                                d = row[0].strip()
                                if d not in distractors and len(distractors) < 4:
                                    distractors.append(d)
                        elif q_word == "When":
                            db_num = con.execute("""
                                SELECT DISTINCT token 
                                FROM corpus 
                                WHERE pos = 'NUM' 
                                  AND token IS NOT NULL 
                                  AND token != ?
                                LIMIT 20
                            """, [correct_text]).fetchall()
                            for row in db_num:
                                d = row[0].strip()
                                if d not in distractors and len(distractors) < 4:
                                    distractors.append(d)
                        con.close()
                    except:
                        pass
                
                # 4. Fallback lists
                fallback_lists = {
                    "PERSON": ["David", "Sarah", "Michael", "Emily", "James", "Jessica"],
                    "GPE": ["Boston", "New York", "Chicago", "San Francisco", "Seattle", "Austin"],
                    "LOC": ["the mountains", "the river", "the valley", "the coast", "the lake"],
                    "FAC": ["the headquarters", "the airport", "the station", "the bridge", "the building"],
                    "DATE": ["yesterday", "next week", "last month", "three days ago", "the following year"],
                    "TIME": ["noon", "midnight", "two hours later", "in the evening", "tomorrow morning"],
                    "ORG": ["the agency", "the association", "the committee", "the department", "the institute"],
                    "PRODUCT": ["the device", "the system", "the software", "the tool", "the equipment"],
                    "EVENT": ["the ceremony", "the conference", "the workshop", "the meeting", "the campaign"],
                    "WORK_OF_ART": ["the report", "the manual", "the guide", "the book", "the document"]
                }
                
                f_list = fallback_lists.get(label, fallback_lists.get("PERSON"))
                for f in f_list:
                    if len(distractors) >= 4:
                        break
                    if f.lower() != correct_text.lower() and f not in distractors:
                        distractors.append(f)
                        
                distractors = distractors[:4]
                all_opts = [correct_text] + distractors
                random.shuffle(all_opts)
                
                letters = ['A', 'B', 'C', 'D', 'E']
                options_dict = {}
                correct_letter = 'A'
                for idx, opt in enumerate(all_opts[:5]):
                    letter = letters[idx]
                    options_dict[letter] = opt
                    if opt == correct_text:
                        correct_letter = letter
                        
                type_5_questions.append({
                    "type": "type_5",
                    "question": q_prompt,
                    "options": options_dict,
                    "correct_option": correct_letter,
                    "correct_text": correct_text,
                    "explanation": f"The blank is correctly completed by '{correct_text}', which is a {category} (labeled as {label}) in the passage."
                })
        except Exception as e:
            print(f"[ERROR] Named Entity Recognition question generation failed: {e}")
            
    random.shuffle(type_5_questions)

    # ----------------------------------------------------
    # ASSEMBLE ALL QUESTIONS (Mix to make exactly 5 items)
    # ----------------------------------------------------
    sentences_shuffled = sentences[:]
    random.shuffle(sentences_shuffled)
    
    def make_type_1(pool):
        if not pool:
            pool.extend(random.sample(sentences, len(sentences)))
        correct_sent = pool.pop(0)
        distractors = []
        explanations = []
        types_to_generate = ["number", "negative", "adverb", "proper_noun"]
        successful_types = []
        
        for dtype in types_to_generate:
            d_info = build_distractor_with_fallback(pool, dtype, successful_types)
            orig_s, mod_s, actual_type = d_info
            if orig_s in pool:
                pool.remove(orig_s)
            distractors.append(mod_s)
            if actual_type != "fallback":
                successful_types.append(actual_type)
            explanations.append(f"Distractor modified via '{actual_type}' rule from sentence: '{orig_s}'")
            
        all_opts = [correct_sent] + distractors
        random.shuffle(all_opts)
        
        letters = ['A', 'B', 'C', 'D', 'E']
        options_dict = {}
        correct_letter = 'A'
        for idx, opt in enumerate(all_opts[:5]):
            letter = letters[idx]
            options_dict[letter] = opt
            if opt == correct_sent:
                correct_letter = letter
                
        explanation_text = "The correct answer is exact from the text: " + f"**{correct_sent}**.\n\n" + "\n".join(explanations)
        return {
            "type": "type_1",
            "question": "Which of the following is mentioned in the text?",
            "options": options_dict,
            "correct_option": correct_letter,
            "correct_text": correct_sent,
            "explanation": explanation_text
        }

    def make_type_2(pool):
        exact_sents = []
        for _ in range(4):
            if not pool:
                pool.extend(random.sample(sentences, len(sentences)))
            exact_sents.append(pool.pop(0))
        d_info = build_distractor_with_fallback(pool, random.choice(["negative", "number", "adverb", "proper_noun"]), ["number", "negative", "adverb", "proper_noun"])
        orig_s, false_sent, actual_type = d_info
        
        all_opts = exact_sents + [false_sent]
        random.shuffle(all_opts)
        
        letters = ['A', 'B', 'C', 'D', 'E']
        options_dict = {}
        correct_letter = 'A'
        for idx, opt in enumerate(all_opts[:5]):
            letter = letters[idx]
            options_dict[letter] = opt
            if opt == false_sent:
                correct_letter = letter
                
        explanation_text = f"The correct answer (not mentioned) is: **{false_sent}**.\nIt is a modified version (using '{actual_type}' rule) of the original sentence: '{orig_s}'."
        return {
            "type": "type_2",
            "question": "Which of the following is NOT mentioned in the text?",
            "options": options_dict,
            "correct_option": correct_letter,
            "correct_text": false_sent,
            "explanation": explanation_text
        }

    # Decide allocations to get 10 questions total with target 2 of each type
    type_1_target = 2
    type_2_target = 2
    type_3_target = 2
    type_4_target = 2 if not type_4_skipped else 0
    type_5_target = 2

    # Add Type 3 if available (up to 2)
    for _ in range(type_3_target):
        if type_3_questions:
            questions.append(type_3_questions.pop(0))
        
    # Add Type 5 if available (up to 2)
    type_5_skipped = len(type_5_questions) == 0
    for _ in range(type_5_target):
        if type_5_questions:
            questions.append(type_5_questions.pop(0))

    # Add Type 4 if available and requested (up to 2: 1 start and 1 end)
    if not type_4_skipped and type_4_questions:
        for q in type_4_questions[:type_4_target]:
            questions.append(q)
            
    # Settle remaining slots with Type 1 & 2 up to 10
    sent_pool = sentences[:]
    random.shuffle(sent_pool)
    
    attempts = 0
    while len(questions) < 10 and attempts < 40:
        attempts += 1
        if not sent_pool:
            sent_pool = sentences[:]
            random.shuffle(sent_pool)
        # Generate Type 1
        try:
            questions.append(make_type_1(sent_pool))
        except:
            pass
        if len(questions) >= 10:
            break
            
        if not sent_pool:
            sent_pool = sentences[:]
            random.shuffle(sent_pool)
        # Generate Type 2
        try:
            questions.append(make_type_2(sent_pool))
        except:
            pass

    return {
        "questions": questions[:10],
        "type_4_skipped": type_4_skipped,
        "type_5_skipped": type_5_skipped
    }

def _generate_paragraph_distractors(remaining_sents, segment_sents, count, all_proper_nouns, word_to_pos, pos_map):
    import random
    distractors = []
    for s in remaining_sents:
        if len(distractors) < count:
            distractors.append(s)
            
    modifier_fns = [
        lambda s: _modify_number(s),
        lambda s: _change_negative_to_positive(s),
        lambda s: _change_degree_adverb(s, word_to_pos, pos_map),
        lambda s: _swap_proper_noun(s, all_proper_nouns, word_to_pos),
        lambda s: _fallback_modify_sentence(s, word_to_pos, pos_map)
    ]
    
    candidate_sources = segment_sents + remaining_sents
    random.shuffle(candidate_sources)
    
    for s in candidate_sources:
        if len(distractors) >= count:
            break
        random.shuffle(modifier_fns)
        for fn in modifier_fns:
            try:
                mod = fn(s)
                if mod and mod != s and mod not in distractors:
                    distractors.append(mod)
                    break
            except:
                pass
                
    while len(distractors) < count:
        distractors.append("An alternate statement detailing related events.")
        
    return distractors[:count]
