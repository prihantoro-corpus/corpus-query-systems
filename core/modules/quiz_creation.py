import duckdb
import pandas as pd
import random
import re
from io import BytesIO
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from collections import Counter

# Set of common punctuation to ignore during content matching
PUNCTUATION = {'.', ',', '!', '?', ';', ':', '(', ')', '[', ']', '"', "'", '“', '”', '‘', '’', '...', '-', '–', '—'}

# Transition/Discourse markers for Section A and B scoring
TRANSITION_MARKERS = {
    'however', 'therefore', 'subsequently', 'consequently', 'furthermore', 
    'moreover', 'thus', 'hence', 'nevertheless', 'meanwhile', 'additionally', 
    'finally', 'instead', 'afterward', 'similarly', 'otherwise', 'accordingly',
    'nonetheless', 'besides'
}

# Multiword markers to check as substrings
MULTIWORD_MARKERS = [
    "these findings", "this result", "these results", "this study", "these data",
    "after that", "as a result", "for instance", "for example", "in addition",
    "on the other hand", "in contrast", "in conclusion", "as well as"
]

# Pronouns/referential continuity items
REFERENTIAL_ITEMS = {
    'he', 'she', 'it', 'they', 'them', 'him', 'her', 'his', 'its', 'their', 
    'this', 'that', 'these', 'those'
}

def get_corpus_sentences(db_path):
    """
    Retrieves all tokens from the corpus database and groups them into sentences.
    Returns a list of dicts: [{'sent_id': int, 'filename': str, 'tokens': [str], 'pos': [str], 'lemmas': [str]}]
    """
    try:
        con = duckdb.connect(db_path, read_only=True)
        # Check if table is empty
        count = con.execute("SELECT count(*) FROM corpus").fetchone()[0]
        if count == 0:
            con.close()
            return []
        
        # Order by token ID to ensure we preserve the sequence
        rows = con.execute("SELECT id, token, pos, lemma, sent_id, filename FROM corpus ORDER BY id").fetchall()
        con.close()
    except Exception as e:
        print(f"Error fetching corpus tokens: {e}")
        return []
    
    sentences = []
    current_sent = None
    
    for row in rows:
        _, token, pos, lemma, sent_id, filename = row
        if current_sent is None or current_sent['sent_id'] != sent_id or current_sent['filename'] != filename:
            if current_sent is not None:
                sentences.append(current_sent)
            current_sent = {
                'sent_id': sent_id,
                'filename': filename,
                'tokens': [],
                'pos': [],
                'lemmas': []
            }
        # Clean verticalised placeholders if present
        token_str = str(token)
        pos_str = str(pos)
        lemma_str = str(lemma)
        current_sent['tokens'].append(token_str)
        current_sent['pos'].append(pos_str)
        current_sent['lemmas'].append(lemma_str)
        
    if current_sent is not None:
        sentences.append(current_sent)
        
    return sentences

def rebuild_sentence_text(tokens):
    """Reconstructs text from a list of tokens with smart spacing before punctuation."""
    text = " ".join(tokens)
    # Remove space before common ending/clause punctuation
    text = re.sub(r'\s+([.,!?;:])', r'\1', text)
    # Handle quotes and parenthesis spacing reasonably
    text = re.sub(r'\(\s+', r'(', text)
    text = re.sub(r'\s+\)', r')', text)
    return text

def get_content_lemmas(sentence_dict):
    """Extracts content lemmas (nouns, verbs, adjectives) to calculate overlap/lexical chains."""
    content_lemmas = set()
    for lemma, pos in zip(sentence_dict['lemmas'], sentence_dict['pos']):
        lemma_low = lemma.lower()
        pos_up = pos.upper()
        if lemma_low in PUNCTUATION or len(lemma_low) <= 1:
            continue
        # Support Penn Treebank (NN*, VB*, JJ*) and UPOS (NOUN, VERB, ADJ, PROPN)
        if (pos_up.startswith('N') or pos_up.startswith('V') or pos_up.startswith('J') or 
                pos_up in {'NOUN', 'VERB', 'ADJ', 'PROPN'}):
            content_lemmas.add(lemma_low)
    return content_lemmas

# =====================================================================
# SECTION A — DISCOURSE COMPLETION GENERATOR
# =====================================================================

def calculate_sentence_removability(sentence, prev_sentence, next_sentence):
    """
    Computes a cohesion/removability score for a candidate sentence.
    Filters out bad candidates (very short, headings, quotes, etc.).
    Returns (is_removable, score)
    """
    tokens = sentence['tokens']
    text = rebuild_sentence_text(tokens).strip()
    
    # 1. Reject criteria
    if len(tokens) < 8 or len(tokens) > 35:
        return False, 0.0
    
    # Heading check: short sentence with no ending punctuation, or all caps
    if len(tokens) < 6:
        return False, 0.0
    if tokens[-1] not in {'.', '?', '!'}:
        return False, 0.0
    if text.isupper():
        return False, 0.0
    
    # Isolated quotation check
    if (text.startswith('"') and text.endswith('"')) or (text.startswith('“') and text.endswith('”')):
        return False, 0.0
        
    score = 0.0
    text_low = text.lower()
    
    # 2. Transition/Discourse Marker Boost
    for marker in TRANSITION_MARKERS:
        # Match word boundaries
        if re.search(r'\b' + re.escape(marker) + r'\b', text_low):
            score += 3.0
            
    for marker in MULTIWORD_MARKERS:
        if marker in text_low:
            score += 4.0
            
    # 3. Referential Pronoun Boost
    for item in REFERENTIAL_ITEMS:
        if re.search(r'\b' + re.escape(item) + r'\b', text_low):
            score += 1.5
            
    # 4. Lexical Overlap / Cosine Similarity with Neighbors
    curr_lemmas = get_content_lemmas(sentence)
    prev_lemmas = get_content_lemmas(prev_sentence) if prev_sentence else set()
    next_lemmas = get_content_lemmas(next_sentence) if next_sentence else set()
    
    overlap_prev = len(curr_lemmas.intersection(prev_lemmas))
    overlap_next = len(curr_lemmas.intersection(next_lemmas))
    
    score += 1.0 * (overlap_prev + overlap_next)
    
    # Sentence is highly coherent and removable if it has some cohesive links
    is_removable = score >= 1.5
    return is_removable, score

def generate_section_a(sentences, num_passages=2):
    """
    Generates Section A items (Discourse Completion).
    Each passage: 20-30 consecutive sentences, 5 sentences removed and replaced by markers.
    Returns: list of passage dicts:
       [{
           'original_passage': str,
           'gapped_passage': str,
           'removed_sentences': [{'marker': str, 'sentence_text': str, 'original_index': int}],
           'options': [{'letter': str, 'sentence_text': str}],
           'correct_mapping': {marker: letter}
       }]
    """
    # Group sentences by filename
    files_map = {}
    for s in sentences:
        files_map.setdefault(s['filename'], []).append(s)
        
    candidate_passages = []
    
    for filename, file_sents in files_map.items():
        if len(file_sents) < 20:
            continue
            
        # We can slide windows of size 20 to 30
        step = 10
        for start_idx in range(0, len(file_sents) - 20, step):
            # Sample passage size between 20 and 30
            p_len = min(25, len(file_sents) - start_idx)
            passage_sents = file_sents[start_idx : start_idx + p_len]
            
            # Filter out if too short overall or average sentence length is too small
            avg_len = sum(len(s['tokens']) for s in passage_sents) / len(passage_sents)
            if avg_len < 8.0:
                continue
                
            # Score all interior sentences for removability (exclude first and last sentences)
            removability_list = []
            for j in range(1, len(passage_sents) - 1):
                curr = passage_sents[j]
                prev = passage_sents[j - 1]
                nxt = passage_sents[j + 1]
                
                is_rem, score = calculate_sentence_removability(curr, prev, nxt)
                if is_rem:
                    removability_list.append((j, score, curr))
                    
            if len(removability_list) < 5:
                continue
                
            # Sort candidate sentences by removability score descending
            removability_list.sort(key=lambda x: x[1], reverse=True)
            
            # Select 5 non-adjacent sentences to remove
            selected_indices = []
            for idx, score, sent in removability_list:
                if len(selected_indices) >= 5:
                    break
                # Ensure no adjacency
                if all(abs(idx - s_idx) > 1 for s_idx in selected_indices):
                    selected_indices.append(idx)
                    
            if len(selected_indices) < 5:
                # Relax adjacency restriction if desperate
                selected_indices = [x[0] for x in removability_list[:5]]
                
            if len(selected_indices) < 5:
                continue
                
            # Sort indices ascending to keep correct order of blanks
            selected_indices.sort()
            
            candidate_passages.append({
                'passage_sents': passage_sents,
                'removed_indices': selected_indices,
                'score': sum(x[1] for x in removability_list[:5])
            })
            
    # Fallback: if no candidates found, relax constraints
    if len(candidate_passages) < num_passages:
        # Let's do a desperate search with relaxed length constraints
        for filename, file_sents in files_map.items():
            if len(file_sents) < 15:
                continue
            for start_idx in range(0, len(file_sents) - 12, 5):
                p_len = min(15, len(file_sents) - start_idx)
                passage_sents = file_sents[start_idx : start_idx + p_len]
                removability_list = []
                for j in range(1, len(passage_sents) - 1):
                    curr = passage_sents[j]
                    if len(curr['tokens']) >= 5:
                        removability_list.append((j, 1.0, curr))
                if len(removability_list) >= 5:
                    selected_indices = [x[0] for x in removability_list[:5]]
                    selected_indices.sort()
                    candidate_passages.append({
                        'passage_sents': passage_sents,
                        'removed_indices': selected_indices,
                        'score': 5.0
                    })
                    
    if not candidate_passages:
        return []
        
    # Shuffle or sort by score and pick top diverse passages
    random.shuffle(candidate_passages)
    candidate_passages.sort(key=lambda x: x['score'], reverse=True)
    
    selected_passages = candidate_passages[:num_passages]
    results = []
    
    for item in selected_passages:
        passage_sents = item['passage_sents']
        removed_indices = item['removed_indices']
        
        # Construct original passage text
        orig_sents = [rebuild_sentence_text(s['tokens']) for s in passage_sents]
        original_passage = " ".join(orig_sents)
        
        # Build gapped passage
        gapped_sents = []
        removed_info = []
        letters = ['A', 'B', 'C', 'D', 'E']
        
        marker_idx = 0
        for idx, sent_text in enumerate(orig_sents):
            if idx in removed_indices:
                marker = f"[MISSING SENTENCE {letters[marker_idx]}]"
                gapped_sents.append(marker)
                removed_info.append({
                    'marker': marker,
                    'sentence_text': sent_text,
                    'original_index': idx
                })
                marker_idx += 1
            else:
                gapped_sents.append(sent_text)
                
        gapped_passage = " ".join(gapped_sents)
        
        # Prepare Shuffled options A-E
        shuffled_sents = removed_info.copy()
        random.shuffle(shuffled_sents)
        
        options = []
        correct_mapping = {}
        for l_idx, letter in enumerate(letters):
            opt_sent = shuffled_sents[l_idx]
            options.append({
                'letter': letter,
                'sentence_text': opt_sent['sentence_text']
            })
            # Find which marker this letter matches
            # The correct mapping connects the marker [MISSING SENTENCE A] to the letter it corresponds to.
            # Let's map marker -> letter
            correct_mapping[opt_sent['marker']] = letter
            
        results.append({
            'original_passage': original_passage,
            'gapped_passage': gapped_passage,
            'removed_sentences': removed_info,
            'options': options,
            'correct_mapping': correct_mapping
        })
        
    return results

# =====================================================================
# SECTION B — SENTENCE REORDERING GENERATOR
# =====================================================================

def calculate_window_cohesion(window_sents):
    """Calculates cohesion score for a 5-sentence consecutive window."""
    score = 0.0
    
    # 1. Avoid windows with very short sentences or headings
    for s in window_sents:
        t_len = len(s['tokens'])
        if t_len < 8 or t_len > 35:
            return 0.0
        # No headings
        text = rebuild_sentence_text(s['tokens'])
        if text.isupper() or s['tokens'][-1] not in {'.', '?', '!'}:
            return 0.0
            
    # 2. Sequential markers & pronouns in sentences 2-5
    for idx, s in enumerate(window_sents[1:], start=1):
        text_low = rebuild_sentence_text(s['tokens']).lower()
        
        # Sequential/ordering markers
        seq_words = {'first', 'second', 'third', 'next', 'then', 'finally', 'subsequently', 'furthermore', 
                     'moreover', 'however', 'therefore', 'consequently', 'thus', 'afterward'}
        for w in seq_words:
            if re.search(r'\b' + re.escape(w) + r'\b', text_low):
                score += 1.5
                
        # Pronouns
        for item in REFERENTIAL_ITEMS:
            if re.search(r'\b' + re.escape(item) + r'\b', text_low):
                score += 1.0
                
    # 3. Lexical overlapping between adjacent sentences
    for idx in range(4):
        lemmas1 = get_content_lemmas(window_sents[idx])
        lemmas2 = get_content_lemmas(window_sents[idx + 1])
        overlap = len(lemmas1.intersection(lemmas2))
        score += 1.0 * overlap
        
    return score

def generate_section_b(sentences, section_a_passages, num_items=5):
    """
    Generates Section B items (Sentence Reordering).
    5 items, each: 5 consecutive coherent sentences extracted from corpus, randomized.
    """
    # Identify sentence ranges that were already used in Section A to prevent duplicates
    used_sent_ids = set()
    for passage in section_a_passages:
        for opt in passage['removed_sentences']:
            used_sent_ids.add(opt['sentence_text'])
            
    # Group sentences by filename
    files_map = {}
    for s in sentences:
        files_map.setdefault(s['filename'], []).append(s)
        
    candidate_windows = []
    
    for filename, file_sents in files_map.items():
        if len(file_sents) < 5:
            continue
            
        for start_idx in range(0, len(file_sents) - 5, 2):
            window = file_sents[start_idx : start_idx + 5]
            
            # Check if any sentence overlaps with Section A text
            texts = [rebuild_sentence_text(s['tokens']) for s in window]
            if any(t in used_sent_ids for t in texts):
                continue
                
            cohesion = calculate_window_cohesion(window)
            if cohesion > 2.0:
                candidate_windows.append({
                    'sentences': window,
                    'cohesion': cohesion,
                    'original_texts': texts
                })
                
    # Fallback if too few candidates
    if len(candidate_windows) < num_items:
        for filename, file_sents in files_map.items():
            if len(file_sents) < 5:
                continue
            for start_idx in range(0, len(file_sents) - 5, 5):
                window = file_sents[start_idx : start_idx + 5]
                texts = [rebuild_sentence_text(s['tokens']) for s in window]
                candidate_windows.append({
                    'sentences': window,
                    'cohesion': 1.0,
                    'original_texts': texts
                })
                
    if not candidate_windows:
        return []
        
    # Sort and pick top unique windows
    random.shuffle(candidate_windows)
    candidate_windows.sort(key=lambda x: x['cohesion'], reverse=True)
    
    selected_windows = []
    for cand in candidate_windows:
        if len(selected_windows) >= num_items:
            break
        # Ensure no overlapping sentences with already selected windows
        overlap = False
        cand_texts = set(cand['original_texts'])
        for sel in selected_windows:
            if cand_texts.intersection(set(sel['original_texts'])):
                overlap = True
                break
        if not overlap:
            selected_windows.append(cand)
            
    # Format Section B items
    results = []
    for item_idx, win in enumerate(selected_windows):
        orig_texts = win['original_texts']
        
        # Prepare randomized order
        indexed_texts = list(enumerate(orig_texts, start=1)) # [(1, sent1), (2, sent2), ...]
        shuffled_indexed = indexed_texts.copy()
        
        # Keep shuffling until the order is actually changed
        attempts = 0
        while len(shuffled_indexed) > 1 and [x[0] for x in shuffled_indexed] == [1,2,3,4,5] and attempts < 10:
            random.shuffle(shuffled_indexed)
            attempts += 1
            
        # The correct order is represented by the indices in their correct positions.
        # E.g., if shuffled is [(3, sent3), (1, sent1), (5, sent5), (2, sent2), (4, sent4)],
        # Then the user sees:
        # 1. sent3
        # 2. sent1
        # 3. sent5
        # 4. sent2
        # 5. sent4
        # To rebuild the sequence 1-2-3-4-5, the correct reordering sequence mapping is:
        # The original sentence 1 is in position 2, original 2 in position 4, original 3 in position 1, original 4 in position 5, original 5 in position 3.
        # So the typed sequence: "3-1-5-2-4" means: original sentence 3 is first, original 1 is second, original 5 is third, etc.
        # Let's double check. If randomized sentences are shown:
        # A = sent3, B = sent1, C = sent5, D = sent2, E = sent4
        # Reordering back: sentence 1 is B, sentence 2 is D, sentence 3 is A, sentence 4 is E, sentence 5 is C.
        # Original order sequence: 3-1-5-2-4 (representing the original index of the sentence displayed at index 0, 1, 2, 3, 4).
        # Yes! The order of original sentence indices is [3, 1, 5, 2, 4].
        correct_seq = "-".join(str(x[0]) for x in shuffled_indexed)
        
        results.append({
            'randomized_sentences': [x[1] for x in shuffled_indexed],
            'original_sentences': orig_texts,
            'correct_sequence': correct_seq,
            'shuffled_mapping': [x[0] for x in shuffled_indexed] # The original index (1-based) of each sentence as displayed
        })
        
    return results

# =====================================================================
# SECTION C — GRAMMAR QUESTION GENERATOR
# =====================================================================

def generate_section_c(sentences, used_texts, num_questions=5):
    """
    Generates Section C (Grammar Questions).
    5 MCQs with option A-E.
    Types: Subject-verb agreement, Verb tense, Articles, Prepositions, Modal verbs, Relative clauses, Passive voice.
    """
    candidates = []
    
    # Supported Grammatical Question Rules
    for s in sentences:
        tokens = s['tokens']
        pos = s['pos']
        lemmas = s['lemmas']
        text = rebuild_sentence_text(tokens)
        
        if text in used_texts:
            continue
        # Try strict criteria first, but fall back to relaxed to ensure we always get questions
        is_strict = True
        if len(tokens) < 10 or len(tokens) > 25:
            is_strict = False
        if not tokens or tokens[-1] not in {'.', '?', '!'}:
            is_strict = False
            
        if not is_strict:
            # Fallback to relaxed criteria: bypass punctuation ending, allow 6-35 tokens
            if len(tokens) < 6 or len(tokens) > 35:
                continue
            
        # 1. PREPOSITIONS
        prepositions = ['in', 'on', 'at', 'for', 'with', 'by', 'about', 'to', 'from', 'of']
        for t_idx, t in enumerate(tokens):
            t_low = t.lower()
            p_up = pos[t_idx].upper()
            # ADP or IN/TO preposition tag
            if t_low in prepositions and (p_up in {'IN', 'TO', 'ADP'}):
                # Generate distractors from other prepositions
                other_preps = [p for p in prepositions if p != t_low]
                distractors = random.sample(other_preps, 4)
                candidates.append({
                    'type': 'Prepositions',
                    'sentence': s,
                    'blank_index': t_idx,
                    'correct_answer': t,
                    'distractors': distractors
                })
                break # Limit 1 question per sentence
                
        # 2. ARTICLES
        articles = {'a', 'an', 'the'}
        for t_idx, t in enumerate(tokens):
            t_low = t.lower()
            if t_low in articles:
                other_arts = list(articles - {t_low})
                distractors = other_arts + ['some', 'any']
                candidates.append({
                    'type': 'Articles',
                    'sentence': s,
                    'blank_index': t_idx,
                    'correct_answer': t,
                    'distractors': distractors
                })
                break
                
        # 3. MODALS
        modals = {'can', 'could', 'will', 'would', 'shall', 'should', 'may', 'might', 'must'}
        for t_idx, t in enumerate(tokens):
            t_low = t.lower()
            p_up = pos[t_idx].upper()
            if t_low in modals and (p_up in {'MD', 'AUX'}):
                other_modals = list(modals - {t_low})
                distractors = random.sample(other_modals, 4)
                candidates.append({
                    'type': 'Modal verbs',
                    'sentence': s,
                    'blank_index': t_idx,
                    'correct_answer': t,
                    'distractors': distractors
                })
                break
                
        # 4. SUBJECT-VERB / VERB TENSE (BE auxiliary)
        be_verbs = {'is', 'are', 'was', 'were'}
        for t_idx, t in enumerate(tokens):
            t_low = t.lower()
            if t_low in be_verbs:
                other_be = list(be_verbs - {t_low})
                distractors = other_be + ['be']
                candidates.append({
                    'type': 'Subject-verb agreement',
                    'sentence': s,
                    'blank_index': t_idx,
                    'correct_answer': t,
                    'distractors': distractors
                })
                break
                
        # 5. RELATIVE CLAUSES
        rel_pronouns = {'who', 'which', 'that', 'whom', 'whose'}
        for t_idx, t in enumerate(tokens):
            t_low = t.lower()
            p_up = pos[t_idx].upper()
            if t_low in rel_pronouns and (p_up in {'WP', 'WDT', 'PRON'}):
                other_rel = list(rel_pronouns - {t_low})
                distractors = random.sample(other_rel, 3) + ['what']
                candidates.append({
                    'type': 'Relative clauses',
                    'sentence': s,
                    'blank_index': t_idx,
                    'correct_answer': t,
                    'distractors': distractors
                })
                break

        # 6. PASSIVE VOICE (Form of BE followed by VBN)
        for t_idx in range(len(tokens) - 1):
            t1 = tokens[t_idx].lower()
            t2 = tokens[t_idx + 1].lower()
            p2 = pos[t_idx + 1].upper()
            l2 = lemmas[t_idx + 1].lower()
            if t1 in {'is', 'are', 'was', 'were', 'been', 'be', 'being'} and (p2 == 'VBN' or p2 == 'VERB'):
                # Blank out the past participle verb
                # Create morphological distractors for l2
                distractors = [l2, l2 + 'ing', l2 + 's', 'to ' + l2]
                # Filter out correct answer in distractors
                distractors = [d for d in distractors if d != t2]
                if len(distractors) < 4:
                    distractors.append(l2 + 'ed')
                distractors = list(set(distractors))[:4]
                while len(distractors) < 4:
                    distractors.append(l2 + 'ment' if not l2.endswith('e') else l2 + 'tion')
                
                candidates.append({
                    'type': 'Passive voice',
                    'sentence': s,
                    'blank_index': t_idx + 1,
                    'correct_answer': tokens[t_idx + 1],
                    'distractors': distractors[:4]
                })
                break

    # Group candidates by type to allow selection of diverse categories
    types_map = {}
    for cand in candidates:
        types_map.setdefault(cand['type'], []).append(cand)
        
    selected_questions = []
    
    # Try to pick 1 from each available type to maximize diversity
    available_types = list(types_map.keys())
    random.shuffle(available_types)
    
    for q_type in available_types:
        if len(selected_questions) >= num_questions:
            break
        type_cands = types_map[q_type]
        # Sort or select randomly
        q = random.choice(type_cands)
        selected_questions.append(q)
        # Record used text to prevent duplication
        used_texts.add(rebuild_sentence_text(q['sentence']['tokens']))
        
    # If still not enough, pick randomly from any type
    if len(selected_questions) < num_questions:
        remaining_candidates = [c for c in candidates if rebuild_sentence_text(c['sentence']['tokens']) not in used_texts]
        random.shuffle(remaining_candidates)
        for q in remaining_candidates:
            if len(selected_questions) >= num_questions:
                break
            selected_questions.append(q)
            used_texts.add(rebuild_sentence_text(q['sentence']['tokens']))
            
    # Format questions MCQ A-E
    results = []
    for q_idx, q in enumerate(selected_questions):
        sentence = q['sentence']
        blank_idx = q['blank_index']
        correct_val = q['correct_answer']
        
        # Build prompt sentence with blank
        tokens_copy = sentence['tokens'].copy()
        tokens_copy[blank_idx] = "_________"
        prompt_text = rebuild_sentence_text(tokens_copy)
        
        # Capitalize start if blank is at index 0
        if blank_idx == 0:
            prompt_text = "_________" + prompt_text[9:]
            
        options = q['distractors'].copy()
        # Add correct answer
        options.append(correct_val)
        options = list(set(options)) # Remove possible duplicates
        while len(options) < 5:
            # Fallback filler
            options.append(correct_val + "_alt")
        random.shuffle(options)
        
        letters = ['A', 'B', 'C', 'D', 'E']
        correct_letter = 'A'
        options_formatted = []
        for l_idx, opt in enumerate(options):
            options_formatted.append({
                'letter': letters[l_idx],
                'text': opt
            })
            if opt == correct_val:
                correct_letter = letters[l_idx]
                
        results.append({
            'question_number': q_idx + 1,
            'type': q['type'],
            'prompt': prompt_text,
            'options': options_formatted,
            'correct_answer': correct_val,
            'correct_letter': correct_letter,
            'original_sentence': rebuild_sentence_text(sentence['tokens'])
        })
        
    return results

# =====================================================================
# SECTION D — MULTIWORD EXPRESSION QUESTION GENERATOR
# =====================================================================

def extract_collocations(db_path):
    """Runs high-performance POS-based collocation query on DuckDB database."""
    try:
        con = duckdb.connect(db_path, read_only=True)
        # Verify POS tags in database to check if they are verticalised empty placeholders or filled
        sample_pos = con.execute("SELECT pos FROM corpus WHERE pos IS NOT NULL LIMIT 5").fetchall()
        pos_is_empty = all(row[0] in {'##', '###', '', None} for row in sample_pos)
        
        if pos_is_empty:
            # POS tagging is not available, we'll fall back to simple bigram collocation extraction
            query = """
            SELECT 
                c1.token AS t1, 
                c2.token AS t2, 
                'NONE' AS p1, 
                'NONE' AS p2, 
                c1.lemma AS l1, 
                c2.lemma AS l2,
                COUNT(*) AS freq
            FROM corpus c1
            JOIN corpus c2 ON c1.id = c2.id - 1 AND c1.sent_id = c2.sent_id
            WHERE length(c1.token) > 2 AND length(c2.token) > 2
            AND c1._token_low NOT IN ('the', 'and', 'for', 'but', 'you', 'that', 'with', 'this', 'have', 'was', 'were')
            AND c2._token_low NOT IN ('the', 'and', 'for', 'but', 'you', 'that', 'with', 'this', 'have', 'was', 'were')
            GROUP BY t1, t2, l1, l2
            ORDER BY freq DESC
            LIMIT 100
            """
        else:
            query = """
            SELECT 
                c1.token AS t1, 
                c2.token AS t2, 
                c1.pos AS p1, 
                c2.pos AS p2, 
                c1.lemma AS l1, 
                c2.lemma AS l2,
                COUNT(*) AS freq
            FROM corpus c1
            JOIN corpus c2 ON c1.id = c2.id - 1 AND c1.sent_id = c2.sent_id
            WHERE (
                -- Verb + Prep
                ((c1.pos LIKE 'V%' OR c1.pos = 'VERB') AND (c2.pos IN ('IN', 'TO', 'ADP'))) OR
                -- Verb + Noun (excluding proper nouns)
                ((c1.pos LIKE 'V%' OR c1.pos = 'VERB') AND (c2.pos = 'NOUN' OR (c2.pos LIKE 'N%' AND c2.pos NOT IN ('NNP', 'NNPS', 'PROPN')))) OR
                -- Adjective + Noun (excluding proper nouns)
                ((c1.pos LIKE 'J%' OR c1.pos = 'ADJ') AND (c2.pos = 'NOUN' OR (c2.pos LIKE 'N%' AND c2.pos NOT IN ('NNP', 'NNPS', 'PROPN')))) OR
                -- Adverb + Adjective
                ((c1.pos LIKE 'R%' OR c1.pos = 'ADV') AND (c2.pos LIKE 'J%' OR c2.pos = 'ADJ')) OR
                -- Noun + Noun (excluding proper nouns)
                ((c1.pos = 'NOUN' OR (c1.pos LIKE 'N%' AND c1.pos NOT IN ('NNP', 'NNPS', 'PROPN'))) AND (c2.pos = 'NOUN' OR (c2.pos LIKE 'N%' AND c2.pos NOT IN ('NNP', 'NNPS', 'PROPN'))))
            )
            AND length(c1.token) > 1 AND length(c2.token) > 1
            AND c1.token NOT IN ('be', 'is', 'are', 'was', 'were', 'have', 'has', 'had', 'do', 'does', 'did')
            GROUP BY t1, t2, p1, p2, l1, l2
            ORDER BY freq DESC
            LIMIT 100
            """
        df_colls = con.execute(query).fetch_df()
        con.close()
        return df_colls
    except Exception as e:
        print(f"Error extracting collocations: {e}")
        return pd.DataFrame()

def is_capitalized(word):
    if not word:
        return False
    w_str = str(word).strip()
    if not w_str:
        return False
    return w_str[0].isupper()

def generate_section_d(db_path, sentences, used_texts, num_questions=5):
    """
    Generates Section D (Multiword Expression Questions).
    5 open-ended questions based on extracted collocations.
    """
    df_colls = extract_collocations(db_path)
    
    if not df_colls.empty:
        # Filter out proper names / capitalized words:
        # Neither t1 nor t2 should start with a capital letter
        df_colls = df_colls[
            df_colls['t1'].apply(lambda x: not is_capitalized(x)) &
            df_colls['t2'].apply(lambda x: not is_capitalized(x))
        ]
        
    if df_colls.empty:
        # Fallback to random common collocations
        common_colls = [
            ("significant", "difference", "Adjective + Noun"),
            ("look", "after", "Verb + Preposition"),
            ("take", "place", "Verb + Noun"),
            ("highly", "successful", "Adverb + Adjective"),
            ("data", "collection", "Noun + Noun")
        ]
        df_colls = pd.DataFrame(common_colls, columns=['t1', 't2', 'pattern'])
        df_colls['freq'] = 10
        
    selected_colls = []
    
    # Shuffle options
    colls_list = df_colls.to_dict('records')
    random.shuffle(colls_list)
    
    # Try to select 5 collocations with diverse patterns
    patterns_used = set()
    for col in colls_list:
        if len(selected_colls) >= num_questions:
            break
            
        # Determine pattern description
        p1 = str(col.get('p1', '')).upper()
        p2 = str(col.get('p2', '')).upper()
        
        pattern_desc = "Collocation"
        if p1.startswith('V') and p2 in {'IN', 'TO', 'ADP'}:
            pattern_desc = "Verb + Preposition"
        elif p1.startswith('V') and (p2.startswith('N') or p2 in {'NOUN', 'PROPN'}):
            pattern_desc = "Verb + Noun"
        elif (p1.startswith('J') or p1 == 'ADJ') and (p2.startswith('N') or p2 in {'NOUN', 'PROPN'}):
            pattern_desc = "Adjective + Noun"
        elif (p1.startswith('R') or p1 == 'ADV') and (p2.startswith('J') or p2 == 'ADJ'):
            pattern_desc = "Adverb + Adjective"
        elif (p1.startswith('N') or p1 in {'NOUN', 'PROPN'}) and (p2.startswith('N') or p2 in {'NOUN', 'PROPN'}):
            pattern_desc = "Noun + Noun"
            
        col['pattern_desc'] = pattern_desc
        
        # We prefer a mix of patterns
        if pattern_desc not in patterns_used or len(patterns_used) >= 4:
            selected_colls.append(col)
            patterns_used.add(pattern_desc)
            
    # Fill up if not enough
    if len(selected_colls) < num_questions:
        for col in colls_list:
            if len(selected_colls) >= num_questions:
                break
            if col not in selected_colls:
                col['pattern_desc'] = col.get('pattern_desc', 'Collocation')
                selected_colls.append(col)
                
    results = []
    for q_idx, col in enumerate(selected_colls):
        t1 = col['t1']
        t2 = col['t2']
        pat = col.get('pattern_desc', 'Collocation')
        
        # Find a sentence in the corpus containing this collocation
        col_sent_text = None
        sentence_obj = None
        matched_i = -1
        context_before = ""
        context_after = ""
        
        for s_idx, s in enumerate(sentences):
            tokens = s['tokens']
            text = rebuild_sentence_text(tokens)
            if text in used_texts:
                continue
            
            # Find consecutive tokens safely
            found = False
            for i in range(len(tokens) - 1):
                tok1 = str(tokens[i])
                tok2 = str(tokens[i+1])
                if tok1.lower() == t1.lower() and tok2.lower() == t2.lower():
                    # Avoid proper names and capitalized words (either token starting with uppercase)
                    t1_cap = tok1 and tok1[0].isupper()
                    t2_cap = tok2 and tok2[0].isupper()
                    if t1_cap or t2_cap:
                        continue
                    matched_i = i
                    found = True
                    break
            
            if found:
                col_sent_text = text
                sentence_obj = s
                
                # Context before (from the same file)
                if s_idx > 0:
                    prev_s = sentences[s_idx - 1]
                    if prev_s['filename'] == s['filename']:
                        context_before = rebuild_sentence_text(prev_s['tokens'])
                
                # Context after (from the same file)
                if s_idx < len(sentences) - 1:
                    next_s = sentences[s_idx + 1]
                    if next_s['filename'] == s['filename']:
                        context_after = rebuild_sentence_text(next_s['tokens'])
                        
                break
                
        if col_sent_text and sentence_obj:
            used_texts.add(col_sent_text)
            
            # Blank out BOTH words of the collocation in the sentence
            tokens_copy = sentence_obj['tokens'].copy()
            if matched_i != -1:
                tokens_copy[matched_i] = "_________"
                tokens_copy[matched_i + 1] = "_________"
                
            main_sentence_blanked = rebuild_sentence_text(tokens_copy)
            
            # Combine context before, main sentence blanked, and context after
            prompt_parts = []
            if context_before:
                prompt_parts.append(context_before)
            prompt_parts.append(main_sentence_blanked)
            if context_after:
                prompt_parts.append(context_after)
                
            prompt_text = " ".join(prompt_parts)
            
            results.append({
                'question_number': q_idx + 1,
                'type': pat,
                'collocation': f"{t1} {t2}",
                'format': 'sentence_cloze',
                'prompt': prompt_text,
                'expected_answer': f"{t1} {t2}",
                'original_sentence': col_sent_text
            })
        else:
            # Fallback to simple completion expression question if sentence not found
            results.append({
                'question_number': q_idx + 1,
                'type': pat,
                'collocation': f"{t1} {t2}",
                'format': 'expression_completion',
                'prompt': f"Complete the collocation: _________ _________ (Hint: {pat})",
                'expected_answer': f"{t1} {t2}",
                'original_sentence': f"Expression: {t1} {t2}"
            })
            
    return results

def generate_section_e(sentences, used_texts, num_questions=5):
    """
    Generates Section E (Sentence Composition Questions).
    5 scramble questions based on sentences from the corpus with 10-15 words.
    """
    candidates = []
    for s in sentences:
        tokens = s['tokens']
        # Filter for sentences with 10 to 15 words
        if not (10 <= len(tokens) <= 15):
            continue
            
        text = rebuild_sentence_text(tokens)
        if text in used_texts:
            continue
            
        # Exclude sentences with too many punctuation mark symbols or non-alphabetic chars
        punc_count = sum(1 for t in tokens if t in {'.', ',', '?', '!', '"', "'", '`', ';', ':', '-', '(', ')', '[', ']'})
        if punc_count > 3:
            continue
            
        # Exclude sentence if it has weird vertical tags
        if any(t in {'##', '###', ''} for t in tokens):
            continue
            
        candidates.append(s)
        
    # Shuffle and select up to 5 items
    random.shuffle(candidates)
    selected_sents = candidates[:num_questions]
    
    # Fallback sentences if corpus is too small or low diversity
    if len(selected_sents) < num_questions:
        fallback_templates = [
            "the researchers conducted a new study on language acquisition",
            "authentic materials are preferred because they avoid artificial structures",
            "students love working with real sentences instead of artificial examples",
            "we hope that more institutions will adopt this corpus driven approach",
            "in conclusion offline corpus query systems are extremely beneficial"
        ]
        for fb_text in fallback_templates:
            if len(selected_sents) >= num_questions:
                break
            # Convert text to tokens
            fb_tokens = fb_text.split()
            selected_sents.append({
                'tokens': fb_tokens,
                'sent_id': 0,
                'filename': 'fallback.txt'
            })
            
    results = []
    for idx, s in enumerate(selected_sents):
        tokens_pool = s['tokens'].copy()
        
        # Strip final sentence punctuation if present so it doesn't float in middle of scramble
        final_punc = ""
        if tokens_pool and tokens_pool[-1] in {'.', '?', '!', ';', ':'}:
            final_punc = tokens_pool.pop()
            
        # Shuffle tokens
        shuffled = tokens_pool.copy()
        random.shuffle(shuffled)
        
        # Format words display list
        shuffled_display = " / ".join(shuffled)
        if final_punc:
            prompt_text = f"Reorganize the following words to form a well-formed sentence (ends with {final_punc}):\n\n{shuffled_display}"
        else:
            prompt_text = f"Reorganize the following words to form a well-formed sentence:\n\n{shuffled_display}"
            
        original_text = rebuild_sentence_text(s['tokens'])
        used_texts.add(original_text)
        
        results.append({
            'question_number': idx + 1,
            'prompt': prompt_text,
            'shuffled_words': shuffled,
            'expected_answer': original_text,
            'original_sentence': original_text
        })
        
    return results

# =====================================================================
# FULL AUTOMATIC OFFLINE QUIZ ENGINE
# =====================================================================

def generate_full_quiz(db_path):
    """
    Generates a full corpus-driven offline language quiz.
    Returns: dict of all sections and details.
    """
    sentences = get_corpus_sentences(db_path)
    if not sentences or len(sentences) < 30:
        return {
            'success': False,
            'error': "The loaded corpus is too small or empty. Please load a larger corpus (at least 30 sentences) to automatically generate quizzes."
        }
        
    # Track used sentences to avoid overlaps between sections
    used_texts = set()
    
    # 1. Section A (Passages)
    section_a = generate_section_a(sentences, num_passages=2)
    for p in section_a:
        used_texts.add(p['original_passage'])
        
    # 2. Section B (Reordering)
    section_b = generate_section_b(sentences, section_a, num_items=5)
    for item in section_b:
        for s in item['original_sentences']:
            used_texts.add(s)
            
    # 3. Section C (Grammar)
    section_c = generate_section_c(sentences, used_texts, num_questions=5)
    for q in section_c:
        used_texts.add(q['original_sentence'])
        
    # 4. Section D (Multiword Expressions)
    section_d = generate_section_d(db_path, sentences, used_texts, num_questions=5)
    for q in section_d:
        used_texts.add(q['original_sentence'])
        
    # 5. Section E (Sentence Composition / Scramble)
    section_e = generate_section_e(sentences, used_texts, num_questions=5)
    
    # Verify we successfully generated all sections
    if not section_a or not section_b or not section_c or not section_d or not section_e:
        return {
            'success': False,
            'error': "Failed to generate all quiz sections due to corpus structural constraints. Try uploading a richer corpus."
        }
        
    return {
        'success': True,
        'section_a': section_a,
        'section_b': section_b,
        'section_c': section_c,
        'section_d': section_d,
        'section_e': section_e
    }

# =====================================================================
# EXPORT TO DOCX GENERATION FUNCTIONS
# =====================================================================

def create_exercises_docx(quiz_data, corpus_name):
    """Generates the Exercise Booklet Word document (.docx) using python-docx."""
    doc = docx.Document()
    
    # Document Styling
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Custom Styles
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("CORTEX Corpus-Driven Language Exercises")
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(20)
    title_run.bold = True
    
    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(f"Data Source: {corpus_name}\nType: Automatically Generated (Offline Mode)")
    sub_run.font.name = 'Times New Roman'
    sub_run.font.size = Pt(11)
    sub_run.italic = True
    doc.add_paragraph() # Spacing
    
    # -------------------------------------------------------------
    # SECTION A
    # -------------------------------------------------------------
    heading_a = doc.add_paragraph()
    h_a_run = heading_a.add_run("SECTION A — Discourse Completion")
    h_a_run.font.size = Pt(14)
    h_a_run.bold = True
    
    inst_a = doc.add_paragraph()
    inst_a_run = inst_a.add_run(
        "Instructions: Read the two passages below. Five sentences have been removed from each passage "
        "and replaced with markers [MISSING SENTENCE A] to [MISSING SENTENCE E]. Match the removed sentences "
        "listed below each passage to their correct locations."
    )
    inst_a_run.italic = True
    
    for idx, passage in enumerate(quiz_data['section_a'], start=1):
        doc.add_paragraph()
        p_title = doc.add_paragraph()
        p_t_run = p_title.add_run(f"Passage {idx}")
        p_t_run.bold = True
        
        pass_p = doc.add_paragraph()
        pass_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pass_p.paragraph_format.line_spacing = 1.15
        pass_p.add_run(passage['gapped_passage'])
        
        doc.add_paragraph("Removed sentences (in shuffled order):")
        for opt in passage['options']:
            opt_p = doc.add_paragraph()
            opt_p.left_indent = Inches(0.5)
            opt_p.add_run(f"{opt['letter']}. ").bold = True
            opt_p.add_run(opt['sentence_text'])
            
    # -------------------------------------------------------------
    # SECTION B
    # -------------------------------------------------------------
    doc.add_page_break()
    heading_b = doc.add_paragraph()
    h_b_run = heading_b.add_run("SECTION B — Sentence Reordering")
    h_b_run.bold = True
    h_b_run.font.size = Pt(14)
    
    inst_b = doc.add_paragraph()
    inst_b_run = inst_b.add_run(
        "Instructions: The sentences in each of the following five items are randomized. "
        "Reorder the sentences in each group into their original coherent sequence. "
        "Write your final answer sequence in the space provided (e.g., 3-1-5-2-4)."
    )
    inst_b_run.italic = True
    
    for idx, item in enumerate(quiz_data['section_b'], start=1):
        doc.add_paragraph()
        p_item = doc.add_paragraph()
        p_i_run = p_item.add_run(f"Item {idx}")
        p_i_run.bold = True
        
        for s_idx, sent in enumerate(item['randomized_sentences'], start=1):
            sent_p = doc.add_paragraph()
            sent_p.left_indent = Inches(0.5)
            sent_p.add_run(f"[{s_idx}] ").bold = True
            sent_p.add_run(sent)
            
        ans_p = doc.add_paragraph()
        ans_p.add_run("Correct sequence order: ________________________").italic = True
        
    # -------------------------------------------------------------
    # SECTION C
    # -------------------------------------------------------------
    doc.add_page_break()
    heading_c = doc.add_paragraph()
    h_c_run = heading_c.add_run("SECTION C — Grammar Questions")
    h_c_run.bold = True
    h_c_run.font.size = Pt(14)
    
    inst_c = doc.add_paragraph()
    inst_c_run = inst_c.add_run(
        "Instructions: Select the grammatically correct option (A-E) to complete each sentence."
    )
    inst_c_run.italic = True
    
    for idx, q in enumerate(quiz_data['section_c'], start=1):
        doc.add_paragraph()
        q_p = doc.add_paragraph()
        q_p.add_run(f"{idx}. ").bold = True
        q_p.add_run(q['prompt'])
        
        for opt in q['options']:
            opt_p = doc.add_paragraph()
            opt_p.left_indent = Inches(0.5)
            opt_p.add_run(f"{opt['letter']}) ").bold = True
            opt_p.add_run(opt['text'])
            
    # -------------------------------------------------------------
    # SECTION D
    # -------------------------------------------------------------
    doc.add_paragraph()
    heading_d = doc.add_paragraph()
    h_d_run = heading_d.add_run("SECTION D — Multiword Expressions")
    h_d_run.bold = True
    h_d_run.font.size = Pt(14)
    
    inst_d = doc.add_paragraph()
    inst_d_run = inst_d.add_run(
        "Instructions: Complete the following sentences or expressions with the appropriate corpus-driven word."
    )
    inst_d_run.italic = True
    
    for idx, q in enumerate(quiz_data['section_d'], start=1):
        doc.add_paragraph()
        q_p = doc.add_paragraph()
        q_p.add_run(f"{idx}. ").bold = True
        q_p.add_run(q['prompt'])
        
    # -------------------------------------------------------------
    # SECTION E
    # -------------------------------------------------------------
    doc.add_page_break()
    heading_e = doc.add_paragraph()
    h_e_run = heading_e.add_run("SECTION E — Sentence Composition")
    h_e_run.bold = True
    h_e_run.font.size = Pt(14)
    
    inst_e = doc.add_paragraph()
    inst_e_run = inst_e.add_run(
        "Instructions: Reorganize the following scrambled words into a grammatically "
        "correct and well-formed sentence."
    )
    inst_e_run.italic = True
    
    for idx, q in enumerate(quiz_data['section_e'], start=1):
        doc.add_paragraph()
        q_p = doc.add_paragraph()
        q_p.add_run(f"{idx}. ").bold = True
        q_p.add_run(q['prompt'])
        
        ans_p = doc.add_paragraph()
        ans_p.add_run("Answer: __________________________________________________________________").italic = True
        
    # Footer info
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.add_run("Generated by CORTEX Corpus Query System © 2026").font.size = Pt(9)
    
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

def create_answer_key_docx(quiz_data, corpus_name):
    """Generates the Answer Key Word document (.docx) using python-docx."""
    doc = docx.Document()
    
    # Styles
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("CORTEX Corpus-Driven Exercises — ANSWER KEY")
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(18)
    title_run.bold = True
    
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(f"Data Source: {corpus_name}")
    sub_run.font.name = 'Times New Roman'
    sub_run.italic = True
    doc.add_paragraph()
    
    # -------------------------------------------------------------
    # SECTION A
    # -------------------------------------------------------------
    heading_a = doc.add_paragraph()
    h_a_run = heading_a.add_run("SECTION A — Discourse Completion Answers")
    h_a_run.bold = True
    h_a_run.font.size = Pt(13)
    
    for idx, passage in enumerate(quiz_data['section_a'], start=1):
        doc.add_paragraph(f"Passage {idx} Mappings:").bold = True
        for marker in ['[MISSING SENTENCE A]', '[MISSING SENTENCE B]', '[MISSING SENTENCE C]', '[MISSING SENTENCE D]', '[MISSING SENTENCE E]']:
            letter = passage['correct_mapping'].get(marker)
            # Find the original sentence text
            original_text = ""
            for o in passage['removed_sentences']:
                if o['marker'] == marker:
                    original_text = o['sentence_text']
                    break
            ans_p = doc.add_paragraph()
            ans_p.left_indent = Inches(0.5)
            ans_p.add_run(f"{marker} ").bold = True
            ans_p.add_run(f"= Option {letter} ").bold = True
            ans_p.add_run(f"({original_text})")
            
    # -------------------------------------------------------------
    # SECTION B
    # -------------------------------------------------------------
    doc.add_paragraph()
    heading_b = doc.add_paragraph()
    h_b_run = heading_b.add_run("SECTION B — Sentence Reordering Answers")
    h_b_run.bold = True
    h_b_run.font.size = Pt(13)
    
    for idx, item in enumerate(quiz_data['section_b'], start=1):
        ans_p = doc.add_paragraph()
        ans_p.add_run(f"Item {idx} Correct Sequence: ").bold = True
        ans_p.add_run(f"{item['correct_sequence']}").bold = True
        
        doc.add_paragraph("Original Order (for reference):")
        for s_idx, sent in enumerate(item['original_sentences'], start=1):
            sent_p = doc.add_paragraph()
            sent_p.left_indent = Inches(0.5)
            sent_p.add_run(f"[{s_idx}] ").bold = True
            sent_p.add_run(sent)
            
    # -------------------------------------------------------------
    # SECTION C
    # -------------------------------------------------------------
    doc.add_paragraph()
    heading_c = doc.add_paragraph()
    h_c_run = heading_c.add_run("SECTION C — Grammar Questions Answers")
    h_c_run.bold = True
    h_c_run.font.size = Pt(13)
    
    for idx, q in enumerate(quiz_data['section_c'], start=1):
        ans_p = doc.add_paragraph()
        ans_p.add_run(f"Question {idx}: ").bold = True
        ans_p.add_run(f"Option {q['correct_letter']} ").bold = True
        ans_p.add_run(f"({q['correct_answer']})").italic = True
        
        ref_p = doc.add_paragraph()
        ref_p.left_indent = Inches(0.5)
        ref_p.add_run("Context: ").italic = True
        ref_p.add_run(q['original_sentence'])
        
    # -------------------------------------------------------------
    # SECTION D
    # -------------------------------------------------------------
    doc.add_paragraph()
    heading_d = doc.add_paragraph()
    h_d_run = heading_d.add_run("SECTION D — Multiword Expressions Answers")
    h_d_run.bold = True
    h_d_run.font.size = Pt(13)
    
    for idx, q in enumerate(quiz_data['section_d'], start=1):
        ans_p = doc.add_paragraph()
        ans_p.add_run(f"Question {idx} ({q['type']}): ").bold = True
        ans_p.add_run(f"{q['expected_answer']}").bold = True
        ans_p.add_run(f"  [Collocation: {q['collocation']}]").italic = True
        
        ref_p = doc.add_paragraph()
        ref_p.left_indent = Inches(0.5)
        ref_p.add_run("Context: ").italic = True
        ref_p.add_run(q['original_sentence'])
        
    # -------------------------------------------------------------
    # SECTION E
    # -------------------------------------------------------------
    doc.add_paragraph()
    heading_e = doc.add_paragraph()
    h_e_run = heading_e.add_run("SECTION E — Sentence Composition Answers")
    h_e_run.bold = True
    h_e_run.font.size = Pt(13)
    
    for idx, q in enumerate(quiz_data['section_e'], start=1):
        ans_p = doc.add_paragraph()
        ans_p.add_run(f"Question {idx}: ").bold = True
        ans_p.add_run(f"{q['expected_answer']}").bold = True
        
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io
